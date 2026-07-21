from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\freertos_study\FreeRTOS今日学习日报-2026-07-15.docx")


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    size = {1: 16, 2: 13, 3: 12}.get(level, 11)
    color = {1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}.get(level, "1F4D78")
    set_font(r, size=size, bold=True, color=color)
    set_spacing(p, before={1: 16, 2: 12, 3: 8}.get(level, 6), after={1: 8, 2: 6, 3: 4}.get(level, 4), line=1.10)


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_label):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    set_spacing(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=10)
            set_spacing(p, after=3, line=1.10)
    fix_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("FreeRTOS 今日学习日报")
    set_font(tr, size=24, bold=True, color="0B2545")
    set_spacing(title, after=3)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("日期：2026年7月15日    代码文件：D:\\freertos_study\\main.c")
    set_font(sr, size=10, color="555555")
    set_spacing(subtitle, after=16)

    add_heading(doc, "一、今日学习主题", 1)
    add_para(doc, "今天围绕 FreeRTOS 的常用内核对象和调度机制进行综合练习。代码不是孤立调用某个 API，而是用 Windows 桌面版 FreeRTOS port 模拟一个“温湿度监控系统”：传感器任务生产数据，控制任务消费数据并判断告警，监督任务汇总事件和通知，统计任务观察系统运行状态，软件定时器周期性产生心跳事件。")
    add_para(doc, "这份练习把前面整理的重点知识从“概念记忆”推进到“对象协作”：任务、队列、互斥量、事件标志组、任务通知、软件定时器、动态内存和栈水位监控都在同一个小系统里形成了闭环。")

    add_heading(doc, "二、今日完成内容", 1)
    add_bullets(doc, [
        "完成了 FreeRTOS 头文件与内核对象的组合使用，代码引入了 task、queue、semphr、event_groups、timers 等模块，说明今天的练习覆盖了多个内核对象。",
        "设计了传感器数据结构和任务参数结构，把队列消息与任务参数从普通变量提升为明确的数据模型。",
        "实现了生产者-消费者流程：传感器任务周期性产生数据并入队，控制任务阻塞等待队列数据并消费。",
        "实现了两类通知机制：事件组用于广播系统事件，任务通知用于点对点告警提醒。",
        "加入了互斥量保护共享输出和共享参数，避免多个任务同时访问同一资源。",
        "加入软件定时器心跳、系统 Tick、队列积压数、剩余堆空间和任务栈水位观察，开始从“能运行”转向“能观察、能定位”。",
    ])

    add_heading(doc, "三、知识点在代码中的应用", 1)
    add_heading(doc, "1. 任务创建、优先级与调度", 2)
    add_para(doc, "今天的任务设计有明显分层：Start 任务负责集中创建业务任务，监督任务优先级最高，传感器任务和控制任务同优先级，统计与参数演示任务优先级较低。这体现了对 FreeRTOS 抢占式调度和优先级设计的理解：关键观察/告警任务要更及时，周期性统计任务不能影响主业务。")
    add_para(doc, "在 Start 任务里批量创建任务时使用临界段保护，创建完成后再删除 Start 任务本身，这对应了任务生命周期管理的实践：初始化型任务完成职责后退出系统调度集合，减少无意义任务占用。")

    add_heading(doc, "2. 队列用于带数据的任务间通信", 2)
    add_para(doc, "队列在本例中承担“传感器样本”的传递。SensorSample_t 把 sequence、temperature、humidity 封装成一条完整消息，xSensorQueue 则把生产者 vSensorTask 和消费者 vControlTask 解耦。这个写法比直接共享全局变量更符合 RTOS 通信习惯：发送方只管投递数据，接收方可阻塞等待，系统不会忙等。")
    add_para(doc, "控制任务使用 portMAX_DELAY 等待队列数据，说明这里主动利用了阻塞态：当没有数据时任务不占 CPU，有数据到达时再由调度器唤醒。")

    add_heading(doc, "3. 事件标志组用于系统事件广播", 2)
    add_para(doc, "事件组被设计成三个 bit：传感器数据就绪、软件定时器心跳、降温请求。它适合表达“某件事发生了”，尤其适合监督任务同时观察多个系统事件。")
    add_para(doc, "vSupervisorTask 通过 xEventGroupWaitBits 同时等待 EVENT_SENSOR_READY 和 EVENT_HEARTBEAT，并设置为任意一个事件到达即可返回；这比为每个事件单独写等待逻辑更清晰，也体现了事件组在多事件同步中的价值。")

    add_heading(doc, "4. 任务通知用于轻量级点对点提醒", 2)
    add_para(doc, "当控制任务检测到温度达到阈值时，除了设置 EVENT_COOLING_REQUEST，还使用 xTaskNotifyGive 直接提醒监督任务。这里把事件组和任务通知做了分工：事件组负责系统状态广播，任务通知负责“控制任务到监督任务”的一对一高温告警。")
    add_para(doc, "监督任务使用 ulTaskNotifyTake 非阻塞读取通知，说明任务通知不仅可以当作简单信号，也可以用于计数式提醒。这个选择比额外创建一个二值信号量更轻量。")

    add_heading(doc, "5. 互斥量保护共享资源", 2)
    add_para(doc, "今天代码中有两处典型共享资源：控制台输出和参数配置结构体。xConsoleMutex 保护 puts/fflush，避免多个任务同时打印造成日志交错；xParamConfigMutex 保护运行期可修改的参数，避免读任务和写任务同时访问同一结构体。")
    add_para(doc, "这个实践把“信号量/互斥量不是为了传数据，而是为了保护资源”的概念落到了代码里。尤其是参数演示任务先加锁、复制快照、再解锁的写法，避免了长时间持锁。")

    add_heading(doc, "6. 软件定时器与事件驱动", 2)
    add_para(doc, "软件定时器被用作 2 秒一次的心跳源，回调函数只做一件短小的事：设置 EVENT_HEARTBEAT。这个设计符合软件定时器回调的使用原则：回调运行在 timer service task 中，应尽量短，不做长时间阻塞操作。")
    add_para(doc, "心跳事件最终被监督任务统一观察，说明软件定时器并不直接承担复杂业务，而是作为周期性事件源接入事件组。")

    add_heading(doc, "7. 系统运行状态与内存/栈观察", 2)
    add_para(doc, "统计任务周期性输出当前 tick、队列积压数量和剩余堆空间，监督任务还读取自身栈高水位。这些内容对应 FreeRTOS 学习中的调试 API：不仅要知道任务在跑，还要能观察系统节拍、队列是否积压、动态内存是否足够、任务栈是否偏小。")
    add_para(doc, "malloc failed hook 和 stack overflow hook 也被保留为异常入口，说明今天已经开始把“错误发生后如何停住并定位”的工程意识加入练习代码。")

    add_heading(doc, "四、代码截图行号索引", 1)
    add_para(doc, "日报正文不直接粘贴代码。需要截图时，可在 D:\\freertos_study\\main.c 中按下面行号范围截取。")
    add_table(doc, ["知识点", "建议截图行号", "截图说明"], [
        ["FreeRTOS 模块引入", "第 4-9 行", "展示 task、queue、semphr、event_groups、timers 等模块同时参与本 demo。"],
        ["事件位定义", "第 40-42 行", "展示事件组中 bit0、bit1、bit2 如何映射系统事件。"],
        ["数据模型与句柄", "第 47-79 行", "展示 SensorSample_t、任务参数结构体，以及队列、互斥量、事件组、定时器、任务句柄。"],
        ["互斥量保护日志", "第 100-104 行", "展示 xSemaphoreTake/xSemaphoreGive 保护共享控制台输出。"],
        ["传感器任务", "第 162-179 行", "展示清除降温请求、发送队列消息、设置传感器事件位、周期延时。"],
        ["控制任务", "第 204-213 行", "展示 xQueueReceive 阻塞接收数据、判断高温、设置事件位并发送任务通知。"],
        ["监督任务", "第 244-270 行", "展示 xEventGroupWaitBits 等待多事件、ulTaskNotifyTake 接收告警、读取栈水位。"],
        ["统计任务", "第 312-317 行", "展示 tick、队列等待数量和剩余堆空间的运行状态观察。"],
        ["带参数任务", "第 327-353 行", "展示 pvParameters、互斥量保护参数快照、按配置周期延时。"],
        ["参数更新任务", "第 364-400 行", "展示运行期修改共享配置，并用同一互斥量保护写入。"],
        ["软件定时器回调", "第 415-422 行", "展示 timer callback 只设置心跳事件位，保持回调短小。"],
        ["任务批量创建", "第 432-486 行", "展示临界段内创建多任务、设置优先级、Start 任务自删除。"],
        ["内核对象创建与启动调度器", "第 501-528 行", "展示互斥量、队列、事件组、软件定时器创建，以及 xTimerStart/vTaskStartScheduler。"],
        ["异常 Hook", "第 543-562 行", "展示内存申请失败和栈溢出时的定位入口。"],
    ], [1.7, 1.35, 3.45])

    add_heading(doc, "五、今日收获", 1)
    add_numbered(doc, [
        "对任务状态的理解更具体了：vTaskDelay、xQueueReceive、xEventGroupWaitBits 都会让任务进入等待/阻塞，系统因此可以把 CPU 让给其他就绪任务。",
        "区分了不同通信工具的职责：队列传递带数据的样本，事件组表达多个系统事件，任务通知负责轻量的一对一告警，互斥量保护共享资源。",
        "理解了任务优先级不是随便设置的：监督任务更高优先级用于及时观察事件，统计任务更低优先级避免干扰主业务。",
        "开始关注工程可观测性：通过 tick、队列积压、free heap、stack high water mark 来判断系统是否健康。",
        "认识到软件定时器回调应保持短小，复杂业务应交给任务处理。"
    ])

    add_heading(doc, "六、问题与后续改进", 1)
    add_bullets(doc, [
        "第 179 行实际延时为 pdMS_TO_TICKS(3000)，但注释写的是 1000ms/1 秒；后续应统一代码与注释，避免复习时误判采样周期。",
        "当前 demo 已经综合使用多个对象，下一步可以重点观察任务优先级变化对日志顺序、队列积压和监督任务响应速度的影响。",
        "可以继续扩展 FromISR 场景，例如模拟中断中发送队列或设置事件位，用来巩固 ISR API 和中断优先级管理。",
        "可以把 vTaskDelay 改成 vTaskDelayUntil 对比周期任务的漂移差异，把时间管理章节的知识再向前推进一步。"
    ])

    add_heading(doc, "七、明日计划", 1)
    add_bullets(doc, [
        "围绕队列阻塞、事件组等待和任务通知分别做一次日志观察，记录任务何时进入阻塞、何时被唤醒。",
        "尝试调整 sensor/control/supervisor 的优先级，观察抢占式调度和时间片轮转在日志中的表现。",
        "检查 FreeRTOSConfig.h 中与任务通知、事件组、软件定时器、栈溢出检测、malloc failed hook 相关的配置项。",
        "补充一版代码注释校准，把实际延时、周期和优先级说明写准确。"
    ])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FreeRTOS 今日学习日报 - 2026年7月15日")
    set_font(fr, size=9, color="777777")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
