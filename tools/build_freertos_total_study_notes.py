from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\freertos_study")
OUT_PATH = ROOT / "FreeRTOS学习总笔记.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "B7C9DD"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, name="Calibri", east_asia="Microsoft YaHei", size=11, color=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_para_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    existing_grid = table._tbl.tblGrid
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        set_style_font(styles[style_name], size=11)
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.25
        styles[style_name].paragraph_format.left_indent = Inches(0.375)
        styles[style_name].paragraph_format.first_line_indent = Inches(-0.188)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=3)
    r = p.add_run("FreeRTOS 学习总笔记")
    set_run_font(r, size=24, bold=True, color=INK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sub, after=18)
    sr = sub.add_run("基于个人笔记、2026-07-15 至 2026-07-20 学习日报、main.c 练习代码与运行日志整理")
    set_run_font(sr, size=10, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(meta, after=18)
    mr = meta.add_run(f"整理时间：{datetime.now().strftime('%Y-%m-%d')}    工程：D:\\freertos_study")
    set_run_font(mr, size=10, color=MUTED)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    set_para_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, after=4)
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_spacing(p, after=4)
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=4)
    tr = p.add_run(title + "：")
    set_run_font(tr, size=10, bold=True, color=INK)
    br = p.add_run(body)
    set_run_font(br, size=10)
    set_table_geometry(table, [6.5])
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=2, line=1.10)
        r = p.add_run(h)
        set_run_font(r, size=9.5, bold=True, color=INK)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, after=2, line=1.10)
            r = p.add_run(str(text))
            set_run_font(r, size=9.2)

    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


KNOWLEDGE_ROWS = [
    ["工程与配置", "FreeRTOSConfig.h、Windows MSVC-MingW port、动态分配、Hook、统计功能", "配置 1ms tick、抢占调度、时间片、128KB heap、事件组/队列集/任务通知/定时器/Tickless 等功能", "工程能启动，日志显示 All Task Created；当前 run_stdout 中 free_heap=60440"],
    ["任务与调度", "任务创建、优先级、阻塞态、挂起态、Start 任务自删除", "创建 supervisor/sensor/control/stats/param/suspend/task-api/delay/notify/low-power 等多任务", "日志显示多任务交替运行；vTaskList 中可见 Running/Blocked/Suspended 等状态"],
    ["队列", "固定长度消息复制、生产者消费者、阻塞等待", "SensorTask 发送 SensorSample_t，ControlTask 阻塞接收", "seq=1/2/3 等样本被 produced 与 consumed；queued=0 表示未积压"],
    ["互斥量", "共享资源保护、优先级继承场景、短临界区", "xConsoleMutex 保护日志，xParamConfigMutex 保护可变参数快照", "多任务日志按行输出，参数任务稳定打印 warn_temp/period_ticks"],
    ["事件组", "bit 位表达多事件、任意位/全部位等待、事件广播", "EVENT_SENSOR_READY、EVENT_HEARTBEAT、EVENT_COOLING_REQUEST", "supervisor 日志观察到 sensor event 与 heartbeat event"],
    ["任务通知", "TCB 内置轻量通信，可计数、置位、传值", "n/w/普通字符三类按键分别触发 ulTaskNotifyTake、xTaskNotifyWait、eSetValueWithOverwrite", "高温告警已有 notification count；第 6 组交互代码已实现，需补充按键运行日志"],
    ["软件定时器", "Timer Service Task、Timer Queue、周期/一次性、短回调", "heartbeat 周期定时器每 2000ms 设置 EVENT_HEARTBEAT", "supervisor 每隔约 2s 观察到 heartbeat event"],
    ["时间管理", "vTaskDelay 相对延时；vTaskDelayUntil 绝对周期延时", "delay-relative 与 delay-absolute 两个任务打印 tick", "相对延时约 1500 tick；绝对延时围绕 2000 tick 周期唤醒"],
    ["信号量与队列集", "二值/计数信号量差异；一个任务等待多个对象", "b/c/q 按键触发二值信号量、计数信号量、队列消息，QueueSetMonitor 统一等待", "日志确认延时实验和菜单入口；b/c/q 交互处理代码已实现，需补充按键运行日志"],
    ["任务 API 与可观测性", "优先级查询/修改、任务状态、栈水位、任务列表、运行统计", "api_worker + api_inspector 周期查询并切换 worker 优先级", "task_count=13/16，prio 1->3->1，stack_free=381，运行统计中 IDLE 约 93%"],
    ["内存管理", "FreeRTOS heap、对象创建成本、free heap、malloc failed hook", "动态创建任务、队列、信号量、事件组、定时器，并持续打印 free heap", "启动后 free_heap 与 stats free_heap 可见；失败 Hook 已接入"],
    ["Tickless 低功耗", "预计空闲 tick、进入前处理、抑制 tick、idle hook", "idle hook 计数，l/L 切换 xLowPowerDemoEnabled，教学版 suppress hook 统计 expected ticks", "第 7 组代码已实现；Windows port 只做教学统计，不代表真实 MCU 电流结果"],
]

KNOWLEDGE_API = {
    "工程与配置": "configUSE_PREEMPTION、configTICK_RATE_HZ、configTOTAL_HEAP_SIZE、configUSE_TIMERS、configUSE_TICKLESS_IDLE、INCLUDE_xxx、configASSERT",
    "任务与调度": "xTaskCreate、xTaskCreateStatic、vTaskDelete、vTaskStartScheduler、vTaskDelay、vTaskDelayUntil、vTaskSuspend、vTaskResume",
    "队列": "xQueueCreate、xQueueSend、xQueueSendToBack、xQueueSendToFront、xQueueReceive、xQueuePeek、uxQueueMessagesWaiting",
    "互斥量": "xSemaphoreCreateMutex、xSemaphoreTake、xSemaphoreGive、xSemaphoreCreateRecursiveMutex、xSemaphoreTakeRecursive、xSemaphoreGiveRecursive",
    "事件组": "xEventGroupCreate、xEventGroupSetBits、xEventGroupClearBits、xEventGroupWaitBits、xEventGroupSync",
    "任务通知": "xTaskNotify、xTaskNotifyGive、ulTaskNotifyTake、xTaskNotifyWait、xTaskNotifyAndQuery、xTaskNotifyFromISR",
    "软件定时器": "xTimerCreate、xTimerStart、xTimerStop、xTimerReset、xTimerChangePeriod、xTimerDelete、xTimerPendFunctionCall",
    "时间管理": "xTaskGetTickCount、pdMS_TO_TICKS、vTaskDelay、vTaskDelayUntil、xTaskAbortDelay",
    "信号量与队列集": "xSemaphoreCreateBinary、xSemaphoreCreateCounting、uxSemaphoreGetCount、xQueueCreateSet、xQueueAddToSet、xQueueSelectFromSet、xQueueRemoveFromSet",
    "任务 API 与可观测性": "uxTaskPriorityGet、vTaskPrioritySet、uxTaskGetNumberOfTasks、uxTaskGetSystemState、vTaskGetInfo、xTaskGetHandle、uxTaskGetStackHighWaterMark、eTaskGetState、vTaskList、vTaskGetRunTimeStats",
    "内存管理": "pvPortMalloc、vPortFree、xPortGetFreeHeapSize、xPortGetMinimumEverFreeHeapSize、vPortDefineHeapRegions、vApplicationMallocFailedHook",
    "Tickless 低功耗": "configUSE_TICKLESS_IDLE、configEXPECTED_IDLE_TIME_BEFORE_SLEEP、configPRE_SLEEP_PROCESSING、configPOST_SLEEP_PROCESSING、portSUPPRESS_TICKS_AND_SLEEP、vApplicationIdleHook",
}


API_QUICK_ROWS = [
    ["配置裁剪", "configUSE_PREEMPTION、configUSE_TIME_SLICING、configTICK_RATE_HZ、configMAX_PRIORITIES、configMINIMAL_STACK_SIZE、configTOTAL_HEAP_SIZE、INCLUDE_xxx", "决定内核功能是否编译、调度行为、Tick 周期、堆大小和可选 API。"],
    ["任务创建/删除", "xTaskCreate、xTaskCreateStatic、vTaskDelete", "动态/静态创建任务，删除任务；删除自身时用 NULL。"],
    ["任务状态/控制", "vTaskSuspend、vTaskResume、xTaskResumeFromISR、eTaskGetState、xTaskGetCurrentTaskHandle、xTaskGetHandle", "挂起/恢复任务，查询任务状态或句柄。"],
    ["任务调度", "vTaskStartScheduler、vTaskEndScheduler、vTaskSuspendAll、xTaskResumeAll、taskYIELD/portYIELD、vTaskSwitchContext", "启动调度器、挂起调度器、主动让出 CPU、执行任务切换。"],
    ["任务调试统计", "uxTaskPriorityGet、vTaskPrioritySet、uxTaskGetNumberOfTasks、uxTaskGetSystemState、vTaskGetInfo、uxTaskGetStackHighWaterMark、vTaskList、vTaskGetRunTimeStats", "观察优先级、任务数量、系统快照、栈水位和运行时间。"],
    ["时间管理", "xTaskGetTickCount、xTaskGetTickCountFromISR、pdMS_TO_TICKS、vTaskDelay、vTaskDelayUntil、xTaskAbortDelay", "读取 tick、毫秒转 tick、相对延时、绝对周期延时、打断阻塞延时。"],
    ["队列", "xQueueCreate、xQueueSend、xQueueSendToBack、xQueueSendToFront、xQueueReceive、xQueuePeek、xQueueOverwrite、uxQueueMessagesWaiting", "固定长度消息通信，支持阻塞等待和队列积压观察。"],
    ["队列 FromISR", "xQueueSendFromISR、xQueueReceiveFromISR、xQueueOverwriteFromISR、xQueuePeekFromISR", "中断中操作队列，配合唤醒标志请求上下文切换。"],
    ["信号量/互斥量", "xSemaphoreCreateBinary、xSemaphoreCreateCounting、xSemaphoreCreateMutex、xSemaphoreTake、xSemaphoreGive、uxSemaphoreGetCount", "同步事件、资源计数、共享资源互斥保护。"],
    ["信号量 FromISR", "xSemaphoreGiveFromISR、xSemaphoreTakeFromISR", "中断中释放或获取信号量；互斥量不能在 ISR 中使用。"],
    ["队列集", "xQueueCreateSet、xQueueAddToSet、xQueueRemoveFromSet、xQueueSelectFromSet", "一个任务同时等待多个队列或信号量成员。"],
    ["事件组", "xEventGroupCreate、xEventGroupSetBits、xEventGroupClearBits、xEventGroupWaitBits、xEventGroupSync", "用 bit 位表达多事件或同步屏障。"],
    ["事件组 FromISR", "xEventGroupSetBitsFromISR、xEventGroupClearBitsFromISR", "中断中设置/清除事件位；实际处理可能委托给 daemon/timer task。"],
    ["任务通知", "xTaskNotify、xTaskNotifyGive、ulTaskNotifyTake、xTaskNotifyWait、xTaskNotifyAndQuery", "轻量一对一通信，可作计数、事件位、邮箱值。"],
    ["任务通知 FromISR", "xTaskNotifyFromISR、xTaskNotifyGiveFromISR、vTaskNotifyGiveFromISR", "中断中给任务发送通知。"],
    ["软件定时器", "xTimerCreate、xTimerStart、xTimerStop、xTimerReset、xTimerChangePeriod、xTimerDelete、xTimerIsTimerActive", "创建、启动、停止、复位、改周期、删除和查询软件定时器。"],
    ["定时器服务任务", "xTimerPendFunctionCall、configTIMER_TASK_PRIORITY、configTIMER_QUEUE_LENGTH、configTIMER_TASK_STACK_DEPTH", "把函数调用挂到 timer service task，配置定时器服务任务资源。"],
    ["临界段", "taskENTER_CRITICAL、taskEXIT_CRITICAL、taskENTER_CRITICAL_FROM_ISR、taskEXIT_CRITICAL_FROM_ISR、portDISABLE_INTERRUPTS、portENABLE_INTERRUPTS", "保护短小共享访问或内核关键结构，区分任务上下文和 ISR 上下文。"],
    ["内核列表", "vListInitialise、vListInitialiseItem、vListInsert、vListInsertEnd、uxListRemove、listGET_OWNER_OF_NEXT_ENTRY", "FreeRTOS 内部组织就绪、阻塞、事件等待、定时器列表。"],
    ["内存管理", "pvPortMalloc、vPortFree、xPortGetFreeHeapSize、xPortGetMinimumEverFreeHeapSize、vPortDefineHeapRegions", "FreeRTOS heap 分配、释放、剩余空间和多段堆配置。"],
    ["Hook/低功耗", "vApplicationIdleHook、vApplicationTickHook、vApplicationMallocFailedHook、vApplicationStackOverflowHook、configPRE_SLEEP_PROCESSING、configPOST_SLEEP_PROCESSING", "空闲、Tick、内存失败、栈溢出和低功耗前后处理入口。"],
]


DETAILS = [
    {
        "title": "工程环境与 FreeRTOSConfig.h",
        "apis": KNOWLEDGE_API["工程与配置"],
        "points": [
            "FreeRTOSConfig.h 是裁剪和配置入口：调度策略、tick 频率、堆大小、可选 API、软件定时器、事件组、队列集、任务通知和 Hook 都在这里打开。",
            "当前工程使用 Windows 桌面 FreeRTOS port，适合在没有开发板时先练 API 和调度行为；迁移 MCU 时需要替换 port、启动文件、链接脚本和硬件初始化。",
            "配置启用了抢占式调度、时间片轮转、1ms tick、动态分配、任务统计、栈溢出检查、malloc failed hook、Tickless 教学路径。",
        ],
        "case": "整理了 include\\FreeRTOSConfig.h，并在 main.c 中动态创建队列、事件组、信号量、软件定时器和多个任务，所有创建结果用 configASSERT 检查。",
        "result": "当前运行日志显示 [0][main] FreeRTOS study demo started 和 [0][start] All Task Created. free_heap=60440，说明配置、对象创建和调度器启动链路可用。",
        "refs": "FreeRTOSConfig.h；main.c:1765-1807；run_stdout.txt",
    },
    {
        "title": "任务创建、优先级与调度",
        "apis": KNOWLEDGE_API["任务与调度"],
        "points": [
            "任务创建的核心是分配 TCB 和栈，然后加入就绪列表；调度器根据优先级和任务状态选择运行任务。",
            "阻塞等待是 RTOS 程序的关键写法：任务无事可做时应等待队列、事件、通知或延时，而不是忙等。",
            "Start 任务只负责集中初始化，创建完所有业务任务后调用 vTaskDelete(NULL) 删除自己。",
        ],
        "case": "在 vStartTask 中集中创建 supervisor、sensor、control、stats、param、suspend、task-api、delay、queue-set、notify、low-power 等任务，并用不同优先级体现任务职责。",
        "result": "日志显示多个任务交替输出；task-api 日志中 task_count=13/16，vTaskList 里能看到 api_inspector Running、多个任务 Blocked、susp_worker Suspended。",
        "refs": "main.c:1566-1750；task_api_demo_run.log",
    },
    {
        "title": "队列：生产者消费者通信",
        "apis": KNOWLEDGE_API["队列"],
        "points": [
            "队列用于任务间传递固定长度消息，发送时复制数据，接收任务可阻塞等待。",
            "队列比共享全局变量更适合跨任务传递完整样本，因为它同时提供缓存、同步和唤醒机制。",
            "队列长度和元素大小决定容量；统计队列等待数量可以观察是否出现积压。",
        ],
        "case": "vSensorTask 周期生成 SensorSample_t，写入 xSensorQueue；vControlTask 使用 xQueueReceive(portMAX_DELAY) 消费样本。",
        "result": "日志出现 [2][sensor] produced seq=1 temperature=26 humidity=51 和 [2][control] consumed seq=1 temperature=26 humidity=51；stats 中 queued=0，说明消费及时。",
        "refs": "main.c:385-461；run_stdout.txt；task_api_demo_run.log",
    },
    {
        "title": "互斥量与共享资源保护",
        "apis": KNOWLEDGE_API["互斥量"],
        "points": [
            "互斥量用于保护共享资源，和普通信号量相比具备优先级继承语义，更适合资源互斥。",
            "锁内工作应尽量短，避免把延时、复杂计算或长时间输出放在持锁区域。",
            "对可变配置建议采用加锁后复制快照，解锁后再使用，减少持锁时间。",
        ],
        "case": "xConsoleMutex 统一保护 puts/fflush，避免多任务日志交错；xParamConfigMutex 保护 xParamDemoConfig，参数任务读取快照，更新任务周期修改配置。",
        "result": "日志按完整行输出，没有明显穿插；param 任务稳定打印 label=param-demo warn_temp=28 humidity_step=5 period_ticks=7000。",
        "refs": "main.c:136-207；main.c:566-655；task_api_demo_run.log",
    },
    {
        "title": "事件组：多事件位同步",
        "apis": KNOWLEDGE_API["事件组"],
        "points": [
            "事件组用一个位图表达多个事件，适合广播状态或让任务等待任意/全部条件。",
            "事件组表达“发生了什么”，不携带复杂数据；如果需要传数据，应配合队列。",
            "等待后是否清除事件位要根据语义选择，避免事件重复触发或被过早清除。",
        ],
        "case": "定义 EVENT_SENSOR_READY、EVENT_HEARTBEAT、EVENT_COOLING_REQUEST；Supervisor 任务等待 sensor 和 heartbeat，Control 任务高温时设置 cooling request。",
        "result": "日志反复出现 [3][supervisor] sensor event observed 和 [3][supervisor] heartbeat event observed；高温时出现 high temperature 后触发后续 cooling applied。",
        "refs": "main.c:36-43；main.c:469-523；task_api_demo_run.log",
    },
    {
        "title": "任务通知：计数、置位和传值",
        "apis": KNOWLEDGE_API["任务通知"],
        "points": [
            "任务通知是每个任务自带的轻量通信机制，内存开销低，适合一对一同步。",
            "xTaskNotifyGive + ulTaskNotifyTake 适合计数信号量式通知；xTaskNotify(..., eSetBits) + xTaskNotifyWait 适合事件位；eSetValueWithOverwrite 可传递 32 位值。",
            "任务通知不是广播机制；如果要多个任务同时感知同一事件，应使用事件组或队列等对象。",
        ],
        "case": "ControlTask 对 Supervisor 使用 xTaskNotifyGive 做高温告警；控制台 n/w/普通字符分别驱动三个通知 demo 任务。",
        "result": "高温样本 seq=3 后日志显示 notification count=1，seq=4 后 notification count=2。第 6 组按键通知代码已完成，但当前保存日志只捕获到菜单入口，需补充一次交互运行记录。",
        "refs": "main.c:431-523；main.c:838-929；main.c:1149-1269；task_api_demo_run.log；run_stdout.txt",
    },
    {
        "title": "软件定时器：周期心跳",
        "apis": KNOWLEDGE_API["软件定时器"],
        "points": [
            "软件定时器由 Timer Service Task 统一管理，启动/停止等命令通过 Timer Queue 异步处理。",
            "回调函数运行在 Timer Service Task 上下文中，应短、快、不阻塞。",
            "周期定时器到期后自动重装，一次性定时器到期后停止。",
        ],
        "case": "xHeartbeatTimer 使用 xTimerCreate 创建，周期 2000ms，pdTRUE 自动重装；vHeartbeatTimerCallback 只设置 EVENT_HEARTBEAT。",
        "result": "日志在约 2s、4s、6s 等时间点反复出现 heartbeat event observed，说明周期定时器和事件组联动正常。",
        "refs": "main.c:955-960；main.c:1780-1784；main.c:1813；task_api_demo_run.log",
    },
    {
        "title": "时间管理：相对延时与绝对周期延时",
        "apis": KNOWLEDGE_API["时间管理"],
        "points": [
            "vTaskDelay 是相对延时，从调用时刻开始计时，适合简单让出 CPU。",
            "vTaskDelayUntil 使用上次唤醒时间作为基准，适合采样、控制、心跳等固定周期任务。",
            "延时会让任务进入阻塞态，到期后回到就绪态，不会忙等占 CPU。",
        ],
        "case": "vRelativeDelayDemoTask 打印 vTaskDelay 前后 tick；vAbsoluteDelayDemoTask 使用 xLastWakeTime 和 2000 tick 周期打印 wait/woke。",
        "result": "numbered_log_check.log 中 relative: tick 56 到 1557 约 1501 tick；absolute: wait-from=46，woke tick=2046，再按 4046 附近继续唤醒，周期稳定在约 2000 tick。",
        "refs": "main.c:978-1041；numbered_log_check.log；delay_sem_queue_set_demo_run.log",
    },
    {
        "title": "信号量与队列集",
        "apis": KNOWLEDGE_API["信号量与队列集"],
        "points": [
            "二值信号量只有空/满，适合一次同步事件；计数信号量可累计 token，适合资源计数或事件累计。",
            "队列集允许一个任务同时等待多个队列或信号量，但返回就绪成员后还必须对具体对象 take/receive。",
            "队列集容量要覆盖所有成员可能产生的事件数量，创建并加入成员的顺序要清楚。",
        ],
        "case": "xDemoQueueSet 同时包含二值信号量、计数信号量和 uint32_t 队列；控制台 b/c/q 按键分别触发三类事件源。",
        "result": "当前日志确认了第 5 组延时任务和菜单入口；b/c/q 交互路径已在代码中实现，但保存日志未捕获实际按键处理结果，后续应补一份按键实验日志。",
        "refs": "main.c:1068-1130；main.c:741-836；main.c:1770-1798；numbered_log_check.log",
    },
    {
        "title": "任务 API、栈水位和运行统计",
        "apis": KNOWLEDGE_API["任务 API 与可观测性"],
        "points": [
            "任务查询 API 是调试工具：可读取优先级、状态、任务数量、栈水位、任务列表和运行时间统计。",
            "vTaskPrioritySet 可运行时调整任务优先级，可能立即影响调度。",
            "vTaskList 和 vTaskGetRunTimeStats 适合学习和诊断，不适合高频放在实时路径。",
        ],
        "case": "api_worker 周期打印自身优先级；api_inspector 按任务名查找 worker，切换优先级并输出 vTaskGetInfo、uxTaskGetSystemState、vTaskList 和 vTaskGetRunTimeStats。",
        "result": "日志显示 prio:1->3 和 prio:3->1，stack_free=381；vTaskList 中列出 13 个任务，运行时间统计显示 IDLE 约 93%，说明大多数时间系统处于空闲/阻塞等待。",
        "refs": "main.c:1382-1561；task_api_demo_run.log",
    },
    {
        "title": "内存管理与异常 Hook",
        "apis": KNOWLEDGE_API["内存管理"],
        "points": [
            "动态创建任务、队列、信号量、事件组、定时器都会消耗 FreeRTOS heap。",
            "xPortGetFreeHeapSize 用于观察剩余堆；malloc failed hook 用于分配失败时停机定位。",
            "stack overflow hook 与栈高水位结合，可帮助判断任务栈大小是否合理。",
        ],
        "case": "配置 128KB FreeRTOS heap；启动后和 stats 任务中打印 free_heap；实现 vApplicationMallocFailedHook 和 vApplicationStackOverflowHook。",
        "result": "当前 run_stdout 显示启动后 free_heap=60440；早期 task_api_demo_run.log 显示 free_heap=22248。不同阶段任务数量和配置不同，但都证明 free heap 已纳入观察。",
        "refs": "include\\FreeRTOSConfig.h；main.c:539-557；main.c:1837-1865；run_stdout.txt；task_api_demo_run.log",
    },
    {
        "title": "Tickless 低功耗学习路径",
        "apis": KNOWLEDGE_API["Tickless 低功耗"],
        "points": [
            "Tickless idle 的目标是在预计空闲时间足够长时抑制周期 tick，让 MCU 进入更长睡眠。",
            "进入低功耗前可通过 hook 决定是否允许睡眠；睡眠后需要补偿 tick 并恢复外设。",
            "Windows 仿真端无法代表真实电流结果，但能练习 expected idle、允许/阻止和统计链路。",
        ],
        "case": "vApplicationIdleHook 统计空闲次数；vLowPowerPreSuppressTicksAndSleep 可把 expected idle time 置 0 来阻止低功耗；vLowPowerSuppressTicksAndSleep 记录进入次数和预计空闲 tick。",
        "result": "第 7 组低功耗日志和 l/L 按键入口已实现；当前保存日志只捕获到菜单入口，真实进入/阻止统计需补充一次更长时间或交互运行记录。",
        "refs": "main.c:1279-1360；main.c:1825-1830；run_stdout.txt",
    },
]


ADDITIONAL_DETAILS = [
    {
        "title": "静态创建与任务生命周期",
        "apis": "xTaskCreateStatic、xTaskCreate、vTaskDelete、configSUPPORT_STATIC_ALLOCATION、configSUPPORT_DYNAMIC_ALLOCATION",
        "points": [
            "动态创建由 FreeRTOS heap 分配 TCB 和任务栈，使用方便，但要持续关注 heap 余量和碎片风险。",
            "静态创建由用户提供 StaticTask_t 和 StackType_t 数组，适合安全等级高、内存布局需要确定或不希望运行期动态分配的场景。",
            "任务函数不应自然返回；删除任务时要理解删除自身、删除其他任务以及空闲任务回收资源之间的差异。",
        ],
        "case": "当前 demo 主要使用 xTaskCreate 做动态创建，并用 vStartTask 完成集中创建后 vTaskDelete(NULL) 删除初始化任务。后续可补一个 xTaskCreateStatic 对照实验，把 TCB 和栈显式放到静态区。",
        "result": "动态创建路径已经验证成功；静态创建在当前代码中尚未实现，适合作为下一步补充练习。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\03-FreeRTOS任务创建和删除.md；main.c:1566-1750",
    },
    {
        "title": "中断管理与 FromISR 规则",
        "apis": "xQueueSendFromISR、xSemaphoreGiveFromISR、xEventGroupSetBitsFromISR、xTaskNotifyFromISR、xTaskResumeFromISR、portYIELD_FROM_ISR",
        "points": [
            "FreeRTOS 规定只有优先级不高于 configMAX_SYSCALL_INTERRUPT_PRIORITY 管理范围的中断，才能调用 FromISR API。",
            "PendSV 和 SysTick 通常配置为最低优先级，PendSV 专门负责延迟上下文切换，避免在普通中断里直接做复杂切换。",
            "FromISR API 常带有 pxHigherPriorityTaskWoken 参数；若唤醒了更高优先级任务，需要请求一次上下文切换。",
        ],
        "case": "当前 Windows demo 还没有真实 ISR 输入，控制台按键临时代替外部事件源。可在后续 MCU/仿真扩展中把按键或定时器中断改成 FromISR 版本，分别发送队列、信号量、事件组和任务通知。",
        "result": "概念已从笔记补入总表；代码层面还属于待补案例，文档中标为后续实验方向。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\05-FreeRTOS中断管理简介.md；04-FreeRTOS的任务挂起与恢复.md",
    },
    {
        "title": "临界段与调度器挂起",
        "apis": "taskENTER_CRITICAL、taskEXIT_CRITICAL、taskENTER_CRITICAL_FROM_ISR、taskEXIT_CRITICAL_FROM_ISR、vTaskSuspendAll、xTaskResumeAll",
        "points": [
            "临界段通过屏蔽可管理中断保护短小关键区，适合保护不能被任务切换或中断打断的共享访问。",
            "vTaskSuspendAll 只是挂起调度器，不关闭中断；ISR 仍然会执行，因此它不能替代中断级临界段。",
            "调度器挂起适合较长但不希望关中断的代码段；恢复时 xTaskResumeAll 可能触发一次任务切换。",
        ],
        "case": "vStartTask 中创建一批任务时使用 taskENTER_CRITICAL/taskEXIT_CRITICAL 包住创建流程，避免初始化过程被中途打断。",
        "result": "当前代码已使用任务级临界段；调度器挂起/恢复还未做单独实验，可补一个保护长列表遍历或批量状态读取的案例。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\06-FreeRTOS临界段代码保护.md；07-任务调度器的挂起和恢复.md；main.c:1572-1749",
    },
    {
        "title": "列表、就绪表和阻塞表",
        "apis": "vListInitialise、vListInitialiseItem、vListInsert、vListInsertEnd、uxListRemove、listGET_OWNER_OF_NEXT_ENTRY",
        "points": [
            "FreeRTOS 内核大量使用双向链表组织任务：就绪列表按优先级分组，阻塞列表按唤醒 tick 排序，事件列表挂等待对象。",
            "TCB 内部通常含有状态列表项和事件列表项，使同一个任务能挂到不同内核列表中。",
            "理解 List_t、ListItem_t、MiniListItem_t，有助于阅读 tasks.c、queue.c、timers.c 的源码。",
        ],
        "case": "当前 demo 没有直接调用 list API，因为它们主要属于内核内部机制；但任务阻塞、事件等待、定时器列表和队列等待都在间接使用这些列表结构。",
        "result": "通过 vTaskList、eTaskGetState、阻塞等待日志，可以从外部观察列表机制的结果：任务在 Running、Blocked、Suspended 之间切换。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\08-FreeRTOS列表和列表项.md；FreeRTOS-KernelV11.3.0\\list.c",
    },
    {
        "title": "调度器启动、首次任务启动与上下文切换",
        "apis": "vTaskStartScheduler、xPortStartScheduler、prvStartFirstTask、vPortSVCHandler、xPortPendSVHandler、vTaskSwitchContext、portYIELD",
        "points": [
            "vTaskStartScheduler 会创建空闲任务和定时器服务任务，并调用移植层启动第一个任务。",
            "Cortex-M 上通常通过 SVC 启动第一个任务，通过 PendSV 保存/恢复上下文，PSP/MSP 分工保证任务栈和异常栈职责清晰。",
            "任务切换本质是保存当前任务寄存器上下文、更新 pxCurrentTCB、恢复下一个任务上下文。",
        ],
        "case": "main() 完成内核对象创建、启动 heartbeat timer 后调用 vTaskStartScheduler；后续所有业务都在任务和定时器回调中运行。",
        "result": "运行日志中 main 只打印启动信息，之后全部为任务日志，说明调度器已接管控制流。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\09-FreeRTOS任务调度.md；main.c:1810-1815",
    },
    {
        "title": "时间片轮询与同优先级任务",
        "apis": "configUSE_TIME_SLICING、configIDLE_SHOULD_YIELD、vTaskDelay、taskYIELD",
        "points": [
            "时间片轮询只发生在同优先级就绪任务之间；高优先级任务就绪时仍然优先抢占。",
            "任务主动调用 vTaskDelay 或等待队列/事件会进入阻塞态，调度器立即选择其他就绪任务。",
            "关闭时间片后，同优先级任务通常需要主动阻塞或 yield，才会给同级任务运行机会。",
        ],
        "case": "sensor 和 control 使用相同优先级，worker/stats/param 等低优先级任务通过延时主动让出 CPU，使日志能看到多任务交替。",
        "result": "日志中多个同/近优先级任务按 tick 和阻塞条件交替输出；任务 API 统计中 IDLE 占比较高，说明任务大多在阻塞等待而非忙等。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\10-FreeRTOS时间片轮询.md；main.c:1583-1605；task_api_demo_run.log",
    },
    {
        "title": "消息队列高级用法与 ISR 通信",
        "apis": "xQueueSendToFront、xQueueSendToBack、xQueuePeek、xQueueOverwrite、xQueueReset、xQueueSendFromISR、xQueueReceiveFromISR",
        "points": [
            "队列不只是普通 send/receive，还支持队头发送、队尾发送、窥视、覆盖写入和复位。",
            "长度为 1 的队列常配合 xQueueOverwrite 做最新值邮箱；多元素队列更适合保留历史消息。",
            "ISR 到任务通信要使用 FromISR 版本，并避免在中断中调用可能阻塞的普通 API。",
        ],
        "case": "当前 xSensorQueue 用最常见的生产者消费者模式传递 SensorSample_t；队列集 demo 中另有 uint32_t 队列成员用于按键消息。",
        "result": "基础队列通信已经跑通；覆盖队列、队头插入和 FromISR 队列通信还可作为后续专项实验。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\13-FreeRTOS消息队列.md；main.c:385-461；main.c:1774",
    },
    {
        "title": "内存算法 heap_1、heap_2、heap_4、heap_5 对比",
        "apis": "pvPortMalloc、vPortFree、xPortGetFreeHeapSize、xPortGetMinimumEverFreeHeapSize、vPortDefineHeapRegions",
        "points": [
            "heap_1 只分配不释放，最简单、确定性好，适合系统启动后不删除对象的场景。",
            "heap_2 支持释放但不合并相邻空闲块，长期运行可能产生碎片。",
            "heap_4 支持释放并合并相邻空闲块，是常见 demo 和中小项目的实用选择。",
            "heap_5 支持多个不连续内存区域，适合 MCU 内部 RAM、外部 RAM 等多段内存统一管理。",
        ],
        "case": "当前工程链接 heap_4.c，并通过 xPortGetFreeHeapSize 观察动态创建对象后的剩余堆。",
        "result": "free_heap 日志能反映对象创建后的资源余量；后续可补 pvPortMalloc/vPortFree 对照实验，观察释放前后和最小剩余堆。",
        "refs": "D:\\freertos_book\\FreeRTOS\\01-笔记\\20-FreeRTOS内存管理.md；FreeRTOS-KernelV11.3.0\\portable\\MemMang\\heap_4.c",
    },
]


def add_detail_section(doc, item):
    add_heading(doc, item["title"], 2)
    add_para(doc, "重点 API：" + item.get("apis", "见本节说明。"), bold_prefix="重点 API：")
    add_para(doc, "重点知识点：", bold_prefix="重点知识点：")
    add_bullets(doc, item["points"])
    add_para(doc, "案例练习：" + item["case"], bold_prefix="案例练习：")
    add_para(doc, "运行结果/观察结论：" + item["result"], bold_prefix="运行结果/观察结论：")
    add_para(doc, "依据：" + item["refs"], bold_prefix="依据：")


def build_doc():
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "总览",
        "这份总笔记把原始笔记中的 FreeRTOS 知识点，与最近几天的日报、main.c 代码实现和日志结果对齐。主线是：先把调度和配置跑起来，再围绕任务、队列、互斥量、事件组、任务通知、定时器、延时、信号量、队列集、内存和低功耗逐步做案例。",
    )

    add_heading(doc, "一、学习主线", 1)
    add_numbers(
        doc,
        [
            "基础框架：搭好 Windows + FreeRTOS Kernel V11.3.0 仿真工程，理解 FreeRTOSConfig.h 与动态对象创建。",
            "对象协作：用温湿度监控 demo 把任务、队列、互斥量、事件组、任务通知和软件定时器连成闭环。",
            "运行控制：加入日志分组、控制台命令、任务挂起恢复、任务 API 查询和运行统计。",
            "专项实验：对比 vTaskDelay/vTaskDelayUntil，练习二值信号量、计数信号量和队列集。",
            "进阶观察：扩展任务通知三种模式、软件定时器内核理解、Tickless 低功耗教学路径和内存/栈诊断。",
        ],
    )

    add_heading(doc, "二、知识点与案例总表", 1)
    enhanced_rows = []
    for row in KNOWLEDGE_ROWS:
        enhanced_rows.append([row[0], row[1], KNOWLEDGE_API.get(row[0], "-"), row[2], row[3]])
    add_table(doc, ["知识点", "重点内容", "重点 API", "你的案例练习", "结果/结论"], enhanced_rows, [0.9, 1.25, 1.55, 1.45, 1.35])

    add_heading(doc, "三、重点 API 速查总表", 1)
    add_table(doc, ["类别", "重点 API / 配置项", "用途"], API_QUICK_ROWS, [1.25, 3.25, 2.0])

    doc.add_page_break()
    add_heading(doc, "四、重点知识点详解", 1)
    for item in DETAILS + ADDITIONAL_DETAILS:
        add_detail_section(doc, item)

    add_heading(doc, "五、当前代码结构索引", 1)
    add_table(
        doc,
        ["模块", "当前 main.c 位置", "说明"],
        [
            ["对象句柄与事件位", "36-125 行", "事件位、日志组、任务句柄、信号量、队列集和低功耗统计变量。"],
            ["日志与控制台", "136-381、714-943 行", "互斥打印、日志分组、键盘命令、交互触发各类 demo。"],
            ["温湿度主流程", "385-523 行", "Sensor/Control/Supervisor 三任务通过队列、事件组和任务通知协作。"],
            ["统计与参数", "539-655 行", "free heap、队列积压、参数任务和参数更新互斥保护。"],
            ["挂起恢复", "656-710 行", "suspend-worker 与 suspend-demo 演示挂起态和恢复。"],
            ["软件定时器", "955-960、1780-1784 行", "heartbeat 周期定时器与短回调。"],
            ["延时与队列集", "978-1130 行", "相对延时、绝对周期延时、队列集监视任务。"],
            ["任务通知", "1149-1269 行", "计数式通知、位通知和值通知三个接收任务。"],
            ["低功耗", "1279-1360、1825-1830 行", "idle hook、Tickless 进入前处理和教学版 suppress hook。"],
            ["任务 API", "1382-1561 行", "优先级、任务状态、栈水位、任务列表、运行时间统计。"],
            ["启动流程", "1566-1815 行", "集中创建任务和内核对象，启动 timer 与 scheduler。"],
            ["异常 Hook", "1837-1865 行", "malloc failed 和 stack overflow 处理入口。"],
        ],
        [1.45, 1.4, 3.65],
    )

    add_heading(doc, "六、运行结果摘录", 1)
    add_table(
        doc,
        ["观察项", "日志摘录/结果", "说明"],
        [
            ["工程启动", "[0][start] All Task Created. free_heap=60440", "当前构建能完成对象与任务创建，并进入调度。"],
            ["队列传递", "produced seq=1 temperature=26 humidity=51；consumed seq=1 temperature=26 humidity=51", "传感器样本从生产者任务传到控制任务。"],
            ["事件组", "sensor event observed；heartbeat event observed", "监督任务能等待并区分多事件位。"],
            ["高温通知", "high temperature, notify supervisor；notification count=1/2", "控制任务用任务通知提醒监督任务。"],
            ["降温请求", "cooling applied seq=4 temperature=29 humidity=49", "控制任务设置请求，传感器任务后续读到事件位并调整湿度。"],
            ["挂起恢复", "vTaskSuspend(suspend-worker): [4] output will stop for 8 seconds", "挂起后 worker 输出暂停，vTaskList 中可见 Suspended。"],
            ["延时对比", "relative 约 1500 tick；absolute 约 2000 tick 周期", "相对延时和绝对周期延时差异可从 tick 中看出。"],
            ["任务 API", "task_count=13；prio:1->3；stack_free=381；IDLE 93%", "任务可观测性和运行统计链路已跑通。"],
            ["待补日志", "b/c/q、n/w/普通字符、l/L", "交互代码已完成，但当前保存日志缺少按键结果，后续建议补一次专项运行。"],
        ],
        [1.3, 3.05, 2.15],
    )

    add_heading(doc, "七、复盘与后续计划", 1)
    add_bullets(
        doc,
        [
            "当前 demo 已经从单点 API 练习升级成一个可观察的小系统：任务负责职责拆分，队列传数据，事件组广播状态，任务通知做点对点提醒，定时器提供周期心跳。",
            "需要修正 main.c 中采样周期注释：代码实际为 pdMS_TO_TICKS(3000)，注释仍写 1000ms/1 秒，容易误导后续复习。",
            "建议补充一份交互日志：依次按 b、c、q、n、w、普通字符、l/L，保存队列集、任务通知和低功耗第 6/7 组输出。",
            "Tickless 目前是 Windows 仿真教学路径，后续迁移到 MCU 时要补外设关断、WFI/WFE、唤醒源和 tick 补偿验证。",
            "内存管理目前主要观察 free heap 和 Hook，后续可以单独增加 pvPortMalloc/vPortFree 小实验，对比 heap_4 释放与合并行为。",
            "建议补充 FromISR 专项：把一个外部输入或模拟中断改成 xQueueSendFromISR、xSemaphoreGiveFromISR、xTaskNotifyFromISR，对照任务上下文 API 的差异。",
            "建议补充静态创建专项：用 xTaskCreateStatic 创建一个低频观察任务，对比动态创建对 heap 的影响。",
        ],
    )

    add_heading(doc, "八、资料来源", 1)
    add_bullets(
        doc,
        [
            "FreeRTOS重点整理.docx：由原始 Markdown 笔记提炼的核心知识点。",
            "FreeRTOS软件定时器内核学习笔记.docx：软件定时器、Timer Queue 和 Timer Service Task 专项笔记。",
            "FreeRTOS今日学习日报-2026-07-15.docx：综合对象协作 demo。",
            "FreeRTOS今日学习日报-2026-07-16.docx：日志分组、挂起恢复、任务 API 与运行统计。",
            "FreeRTOS今日学习日报-2026-07-17.docx：时间管理、信号量和队列集。",
            "FreeRTOS今日学习日报-2026-07-20.docx：任务通知、软件定时器、Tickless 和内存观察。",
            "main.c、include\\FreeRTOSConfig.h、run_stdout.txt、task_api_demo_run.log、delay_sem_queue_set_demo_run.log、numbered_log_check.log：代码与运行结果依据。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("FreeRTOS 学习总笔记")
    set_run_font(r, size=9, color="777777")

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
