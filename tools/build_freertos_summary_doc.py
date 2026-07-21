from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_DIR = Path(r"D:\freertos_book\FreeRTOS\01-笔记")
OUT_PATH = Path(r"D:\freertos_study\FreeRTOS重点整理.docx")


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run(text)
    r.text = text
    colors = {1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}
    sizes = {1: 16, 2: 13, 3: 12}
    set_run_font(r, size=sizes.get(level, 11), bold=True, color=colors.get(level, "1F4D78"))
    set_paragraph_spacing(p, before={1: 18, 2: 14, 3: 10}.get(level, 8), after={1: 10, 2: 7, 3: 5}.get(level, 4))
    return p


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    set_paragraph_spacing(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)
        set_paragraph_spacing(p, after=4)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)
        set_paragraph_spacing(p, after=4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=10)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def extract_source_stats():
    md_files = sorted(SOURCE_DIR.glob("*.md"))
    stats = []
    for p in md_files:
        text = p.read_text(encoding="utf-8")
        title = next((line.lstrip("# ").strip() for line in text.splitlines() if line.startswith("# ")), p.stem)
        headings = [line.strip("# ").strip() for line in text.splitlines() if line.startswith("## ")]
        apis = sorted(set(re.findall(r"\b(?:xTask|vTask|uxTask|eTask|xQueue|vQueue|xSemaphore|xEventGroup|xTimer|pvPort|vPort|xPort)[A-Za-z0-9_]*\b", text)))
        stats.append((p.name, title, len(headings), ", ".join(apis[:8])))
    return stats


CHAPTERS = [
    ("基础知识", [
        "任务调度器负责在多个任务之间选择当前要运行的任务，核心依据是任务状态、优先级和调度策略。",
        "FreeRTOS 常见调度方式包括抢占式调度、时间片调度和协程式调度；实际项目以抢占式和同优先级时间片轮转为主。",
        "任务状态重点掌握运行态、就绪态、阻塞态、挂起态：只有就绪态任务才可能被调度为运行态。",
        "就绪列表按优先级组织，阻塞列表管理延时或等待事件的任务，挂起列表管理被显式暂停的任务。"
    ], [
        ("抢占式调度", "高优先级任务就绪后可立即抢占低优先级任务。"),
        ("时间片调度", "相同优先级任务按 Tick 节拍轮流运行，未用完的时间片不会结转。"),
        ("阻塞态", "任务等待延时、队列、信号量、事件等条件，条件满足或超时后回到就绪态。"),
        ("挂起态", "由 vTaskSuspend() 进入，需要 vTaskResume() 或 ISR 恢复接口唤醒。")
    ]),
    ("系统配置文件", [
        "FreeRTOSConfig.h 用来裁剪内核功能、配置调度行为、使能可选 API，并绑定移植层所需的中断宏。",
        "INCLUDE_ 开头的宏通常控制某个 API 是否编译；config 开头的宏通常控制内核功能和资源参数。",
        "配置文件是移植和调试的入口，修改调度、内存、信号量、软件定时器、Tickless 等功能前应先确认相关宏。"
    ], [
        ("INCLUDE_xxx", "决定某些可选 API 是否可用，例如任务删除、挂起、恢复等。"),
        ("configxxx", "配置功能开关、优先级数量、Tick 频率、堆大小、调度策略等。"),
        ("中断相关宏", "用于映射 PendSV、SVC、SysTick 等移植层接口。")
    ]),
    ("任务创建和删除", [
        "任务创建的本质是分配并初始化 TCB 与任务栈，然后把任务加入就绪列表。",
        "动态创建由 FreeRTOS 堆分配 TCB 和栈，使用简单；静态创建由用户提供内存，更适合对内存位置和确定性有要求的场景。",
        "任务删除会把任务从状态列表或事件列表移除；删除自身时，动态分配的内存通常由空闲任务回收。",
        "任务函数不应自然返回，LR 通常会指向错误处理函数，避免任务跑飞。"
    ], [
        ("xTaskCreate()", "动态创建任务，需 configSUPPORT_DYNAMIC_ALLOCATION。"),
        ("xTaskCreateStatic()", "静态创建任务，需用户提供 TCB 和栈空间。"),
        ("vTaskDelete()", "删除任务；参数为 NULL 时删除当前任务。"),
        ("TCB", "保存任务状态、栈指针、优先级、列表项等调度所需信息。")
    ]),
    ("任务挂起与恢复", [
        "挂起是可恢复的暂停，不等同于删除；被挂起任务不会参与调度，直到被恢复。",
        "vTaskSuspend() 会把任务从原列表移除并插入挂起列表；如果挂起的是当前任务，会触发重新调度。",
        "任务中恢复使用 vTaskResume()；中断中恢复使用 xTaskResumeFromISR()，并根据返回值决定是否需要切换上下文。",
        "同一任务即使被挂起多次，通常恢复一次即可重新进入就绪态。"
    ], [
        ("vTaskSuspend()", "挂起指定任务，NULL 表示挂起自身。"),
        ("vTaskResume()", "任务上下文中恢复被挂起任务。"),
        ("xTaskResumeFromISR()", "ISR 中恢复任务，需注意中断优先级受 FreeRTOS 管理。")
    ]),
    ("中断管理", [
        "中断会打断当前程序并进入 ISR，FreeRTOS 通过中断优先级规则决定哪些 ISR 可以调用内核 API。",
        "Cortex-M 的 PendSV 和 SysTick 通常被配置为较低优先级，用于安全地完成任务切换和系统节拍。",
        "调用 FromISR API 的中断优先级不能高于 FreeRTOS 可管理的最高中断优先级，否则可能破坏内核临界区。",
        "理解抢占优先级、子优先级、BASEPRI 屏蔽机制，是排查 ISR 与 RTOS 交互问题的关键。"
    ], [
        ("SysTick", "提供系统 Tick，驱动延时、时间片、超时检查。"),
        ("PendSV", "承担延迟上下文切换，避免在任意 ISR 中直接切换任务。"),
        ("BASEPRI", "屏蔽一部分可管理中断，保护内核临界段。")
    ]),
    ("临界段代码保护", [
        "临界段用于保护不希望被打断的共享资源访问或内核数据结构操作。",
        "任务级临界段通常通过 taskENTER_CRITICAL()/taskEXIT_CRITICAL() 进入和退出。",
        "ISR 中使用专门的 FromISR 临界段接口，保存并恢复 BASEPRI，避免错误地影响更高优先级中断。",
        "临界段应尽量短，不能在临界段内做耗时或可能阻塞的操作。"
    ], [
        ("taskENTER_CRITICAL()", "任务中进入临界段。"),
        ("taskEXIT_CRITICAL()", "任务中退出临界段。"),
        ("taskENTER_CRITICAL_FROM_ISR()", "ISR 中进入临界段并保存屏蔽状态。"),
        ("taskEXIT_CRITICAL_FROM_ISR()", "ISR 中恢复进入前的屏蔽状态。")
    ]),
    ("调度器挂起和恢复", [
        "vTaskSuspendAll() 只是挂起调度器，并不关闭中断；ISR 仍可运行。",
        "调度器挂起期间任务不会切换，但 Tick 和事件可能被挂起记录，恢复时统一处理。",
        "xTaskResumeAll() 恢复调度器，并可能触发一次上下文切换。",
        "它适合保护较长但不希望禁中断的区域；不能替代互斥量来长期保护共享资源。"
    ], [
        ("vTaskSuspendAll()", "挂起任务调度。"),
        ("xTaskResumeAll()", "恢复任务调度，返回值可反映是否已发生切换。")
    ]),
    ("列表和列表项", [
        "FreeRTOS 大量使用双向链表组织任务、延时、事件等待等内核对象。",
        "列表项嵌入在 TCB 或对象结构体中，使任务可以被快速插入、移除和排序。",
        "迷你列表项常作为列表尾节点，简化链表边界处理。",
        "理解 List_t、ListItem_t、MiniListItem_t 能帮助阅读调度器、队列和定时器源码。"
    ], [
        ("List_t", "列表头，记录列表项数量、索引和尾节点。"),
        ("ListItem_t", "普通列表项，携带排序值和所属对象指针。"),
        ("MiniListItem_t", "简化列表项，常用于列表尾节点。")
    ]),
    ("任务调度与上下文切换", [
        "vTaskStartScheduler() 启动调度器，完成硬件配置、创建空闲任务/定时器任务并启动第一个任务。",
        "xPortStartScheduler() 与具体架构相关，通常配置 SysTick、PendSV、SVC 优先级和中断。",
        "首次启动任务需要恢复任务栈中预先构造的寄存器上下文，使 PC 指向任务入口。",
        "PendSV 处理中保存当前任务上下文、调用 vTaskSwitchContext() 选择下一个任务，再恢复新任务上下文。"
    ], [
        ("vTaskStartScheduler()", "启动 FreeRTOS 调度器。"),
        ("xPortStartScheduler()", "移植层启动调度所需硬件。"),
        ("prvStartFirstTask()", "启动第一个任务。"),
        ("vTaskSwitchContext()", "选择下一个要运行的任务。")
    ]),
    ("时间片轮询", [
        "时间片轮询发生在相同优先级任务之间，依赖系统 Tick。",
        "高优先级任务就绪时仍然优先抢占，同优先级任务才轮流运行。",
        "若某任务在时间片内主动阻塞，调度器会立即切到同优先级或更低优先级的下一个可运行任务。",
        "关闭时间片后，同优先级任务通常需要主动让出 CPU 或进入阻塞才能切换。"
    ], [
        ("configUSE_TIME_SLICING", "控制同优先级时间片轮转。"),
        ("SysTick 周期", "决定一个时间片的基本节拍。")
    ]),
    ("任务相关 API", [
        "任务查询 API 用于调试和运行时监控，例如查看优先级、任务数量、任务状态、栈剩余水位等。",
        "vTaskPrioritySet() 可动态修改任务优先级，可能立即导致调度变化。",
        "uxTaskGetStackHighWaterMark() 是排查栈空间是否不足的重要接口。",
        "vTaskList() 和 vTaskGetRunTimeStats() 适合调试，不宜在强实时路径频繁调用。"
    ], [
        ("uxTaskPriorityGet()", "获取任务优先级。"),
        ("vTaskPrioritySet()", "修改任务优先级。"),
        ("uxTaskGetNumberOfTasks()", "获取系统任务数量。"),
        ("uxTaskGetSystemState()", "获取任务状态数组。"),
        ("eTaskGetState()", "查询指定任务状态。"),
        ("uxTaskGetStackHighWaterMark()", "获取栈历史最小剩余。")
    ]),
    ("时间管理", [
        "FreeRTOS 延时基于 Tick，常见接口为 vTaskDelay() 和 vTaskDelayUntil()。",
        "vTaskDelay() 是相对延时，以上一次调用时刻为起点，任务执行时间会影响周期稳定性。",
        "vTaskDelayUntil() 是绝对延时，适合周期任务，可减少任务执行时间造成的漂移。",
        "延时会让任务进入阻塞态，到期后重新进入就绪态，而不是忙等占用 CPU。"
    ], [
        ("vTaskDelay()", "相对延时。"),
        ("vTaskDelayUntil()", "绝对延时，适合周期控制。"),
        ("TickType_t", "系统节拍计数类型。")
    ]),
    ("消息队列", [
        "队列用于任务间或 ISR 到任务之间传递固定长度消息，复制数据而不是共享同一块变量。",
        "队列由队列长度和队列项大小决定容量，满队列写入或空队列读取可按阻塞时间等待。",
        "多个任务等待同一队列时，优先级高的任务优先解除阻塞；同优先级通常按等待顺序处理。",
        "从 ISR 中操作队列应使用 FromISR 版本，并根据唤醒标志请求上下文切换。"
    ], [
        ("xQueueCreate()", "创建队列。"),
        ("xQueueSend()/xQueueSendToBack()", "写入队尾。"),
        ("xQueueSendToFront()", "写入队头。"),
        ("xQueueReceive()", "读取并移除队列项。"),
        ("xQueuePeek()", "读取但不移除。"),
        ("xQueueOverwrite()", "覆盖写入，常用于长度为 1 的队列。")
    ]),
    ("信号量", [
        "信号量本质上基于队列机制，但通常不传递实际数据，只表达资源计数或同步事件。",
        "二值信号量常用于任务同步，计数型信号量适合资源计数或事件累计。",
        "互斥信号量用于共享资源互斥访问，具备优先级继承机制，可缓解优先级翻转。",
        "普通信号量没有优先级继承，不能随意替代互斥量保护共享资源。"
    ], [
        ("xSemaphoreCreateBinary()", "创建二值信号量。"),
        ("xSemaphoreCreateCounting()", "创建计数型信号量。"),
        ("xSemaphoreCreateMutex()", "创建互斥量。"),
        ("xSemaphoreTake()", "获取信号量或互斥量。"),
        ("xSemaphoreGive()", "释放信号量或互斥量。"),
        ("uxSemaphoreGetCount()", "读取当前计数。")
    ]),
    ("队列集", [
        "队列集允许一个任务同时等待多个队列或信号量中的任意一个变为可用。",
        "使用流程通常是创建队列集、把队列/信号量加入集合、等待集合返回就绪成员、再从成员对象读取。",
        "队列集容量要能容纳所有成员可能产生的事件数量，否则会丢失通知或无法加入。",
        "队列集适合多输入等待场景，但会增加理解和调试成本，简单场景优先使用直接队列或任务通知。"
    ], [
        ("xQueueCreateSet()", "创建队列集。"),
        ("xQueueAddToSet()", "添加队列或信号量到集合。"),
        ("xQueueRemoveFromSet()", "从集合移除成员。"),
        ("xQueueSelectFromSet()", "等待并返回有事件的成员。")
    ]),
    ("事件标志组", [
        "事件标志组用一个位图表达多个事件条件，一个任务可以等待任意位或全部位。",
        "configUSE_16_BIT_TICKS 会影响可用事件位数量：常见 32 位 Tick 下可用 24 个事件位。",
        "等待接口可选择退出时清除事件位，也可选择等待任意位或全部位满足。",
        "事件标志组适合多事件同步；若需要传递数据，仍应使用队列或其他通信对象。"
    ], [
        ("xEventGroupCreate()", "创建事件标志组。"),
        ("xEventGroupSetBits()", "设置事件位。"),
        ("xEventGroupClearBits()", "清除事件位。"),
        ("xEventGroupWaitBits()", "等待事件位。"),
        ("xEventGroupSync()", "任务间同步屏障。")
    ]),
    ("任务通知", [
        "任务通知是每个任务自带的轻量级通信机制，可作为二值信号量、计数信号量、事件位或邮箱使用。",
        "优势是速度快、占用内存少；限制是通知目标必须是某个任务，且默认通知槽数量有限。",
        "通知值更新方式包括不更新、按位设置、递增、覆盖写入、不覆盖写入。",
        "接收端关注两个状态：是否正在等待通知，以及通知值如何在进入/退出等待时清除。"
    ], [
        ("xTaskNotify()", "发送任务通知。"),
        ("xTaskNotifyGive()", "以计数方式通知任务。"),
        ("xTaskNotifyWait()", "等待通知并读取通知值。"),
        ("ulTaskNotifyTake()", "以信号量方式获取通知。"),
        ("eNotifyAction", "决定通知值更新方式。")
    ]),
    ("软件定时器", [
        "软件定时器由 FreeRTOS 定时器服务任务统一管理，回调函数在服务任务上下文中执行。",
        "定时器 API 通常把命令发送到定时器命令队列，由服务任务异步处理。",
        "单次定时器超时后停止；周期定时器回调后会自动重新装载。",
        "定时器回调中不能执行会长时间阻塞的操作，否则会影响所有软件定时器命令处理。"
    ], [
        ("xTimerCreate()", "创建软件定时器。"),
        ("xTimerStart()", "启动定时器。"),
        ("xTimerStop()", "停止定时器。"),
        ("xTimerReset()", "复位并重新开始计时。"),
        ("xTimerChangePeriod()", "修改定时器周期。")
    ]),
    ("Tickless 低功耗", [
        "Tickless 通过在空闲时间抑制周期性 Tick 中断，让 MCU 更长时间停留在低功耗模式。",
        "系统会估算下一次必须唤醒的时间，进入睡眠前调整定时器，醒来后补偿丢失的 Tick。",
        "进入 Tickless 的前提包括启用 configUSE_TICKLESS_IDLE、空闲任务正在运行、预计空闲时间达到阈值。",
        "可通过 configPRE_SLEEP_PROCESSING() 和 configPOST_SLEEP_PROCESSING() 在睡眠前后关闭/恢复外设时钟。"
    ], [
        ("configUSE_TICKLESS_IDLE", "使能 Tickless 空闲低功耗。"),
        ("configEXPECTED_IDLE_TIME_BEFORE_SLEEP", "进入睡眠前的最小预计空闲 Tick 数。"),
        ("configPRE_SLEEP_PROCESSING()", "睡眠前用户钩子。"),
        ("configPOST_SLEEP_PROCESSING()", "睡眠后用户钩子。")
    ]),
    ("内存管理", [
        "FreeRTOS 为动态创建任务、队列、信号量、定时器等对象提供自己的堆管理方案。",
        "不直接依赖标准 C 库 malloc/free 的原因包括代码体积、线程安全、时间确定性和碎片控制。",
        "heap_1 只分配不释放；heap_2 支持释放但不合并相邻空闲块；heap_4 支持释放并合并相邻空闲块；heap_5 支持多个非连续内存区域。",
        "实际项目应根据对象生命周期、碎片风险、内存布局选择堆实现，常见示例多使用 heap_4。"
    ], [
        ("pvPortMalloc()", "从 FreeRTOS 堆申请内存。"),
        ("vPortFree()", "释放动态内存。"),
        ("xPortGetFreeHeapSize()", "获取当前剩余堆空间。"),
        ("vPortDefineHeapRegions()", "为 heap_5 指定多个内存区域。")
    ]),
]


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("FreeRTOS 笔记重点整理")
    set_run_font(r, size=24, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("基于 D:\\freertos_book\\FreeRTOS\\01-笔记 中 20 篇 Markdown 笔记提炼")
    set_run_font(sr, size=10, color="555555")
    set_paragraph_spacing(sub, after=18)

    add_heading(doc, "阅读主线", 1)
    add_bullets(doc, [
        "先抓住调度模型：任务状态、优先级、就绪/阻塞/挂起列表，以及 SysTick 和 PendSV 的分工。",
        "再掌握对象模型：任务、队列、信号量、事件组、任务通知、软件定时器都围绕内核对象和等待列表运转。",
        "最后关注工程配置：FreeRTOSConfig.h、内存管理方案、中断优先级、Tickless 低功耗和调试 API。",
    ])

    add_heading(doc, "核心对比速查", 1)
    add_table(doc, ["主题", "重点结论", "常见误区"], [
        ["队列 vs 全局变量", "队列复制消息并提供阻塞/唤醒机制，适合任务间通信。", "全局变量没有天然同步，容易被并发访问打断。"],
        ["二值信号量 vs 互斥量", "二值信号量偏同步，互斥量偏资源保护。", "互斥量有优先级继承，普通信号量没有。"],
        ["事件组 vs 任务通知", "事件组适合多任务多事件位，任务通知适合点对点轻量通信。", "任务通知不能广播给多个任务。"],
        ["临界段 vs 调度器挂起", "临界段屏蔽可管理中断，调度器挂起不关闭中断。", "把调度器挂起当作中断保护会留下竞态。"],
        ["vTaskDelay vs vTaskDelayUntil", "前者相对延时，后者绝对周期延时。", "周期任务用相对延时会累积执行时间漂移。"],
    ], [1.55, 2.8, 2.15])

    add_heading(doc, "源文件清单", 1)
    rows = []
    for name, title_text, heading_count, apis in extract_source_stats():
        rows.append([name, title_text, str(heading_count), apis or "-"])
    add_table(doc, ["文件", "主题", "二级章节数", "识别到的部分 API"], rows, [2.0, 1.7, 0.8, 2.0])

    doc.add_page_break()
    add_heading(doc, "分章节重点", 1)
    for idx, (title_text, points, api_rows) in enumerate(CHAPTERS, start=1):
        add_heading(doc, f"{idx:02d}. {title_text}", 2)
        add_bullets(doc, points)
        if api_rows:
            add_table(doc, ["关键词 / API", "掌握要点"], api_rows, [2.0, 4.3])

    doc.add_page_break()
    add_heading(doc, "复习检查清单", 1)
    add_numbered(doc, [
        "能画出任务状态转换关系，并说明就绪、阻塞、挂起之间的区别。",
        "能解释抢占式调度、时间片调度、PendSV 上下文切换和 SysTick 节拍的关系。",
        "能根据场景选择队列、信号量、互斥量、事件组、任务通知或软件定时器。",
        "能说明哪些 API 可以在 ISR 中使用，以及中断优先级为什么必须受 FreeRTOS 规则约束。",
        "能根据内存生命周期选择 heap_1、heap_2、heap_4 或 heap_5，并知道动态对象删除后的内存回收路径。",
        "能配置 Tickless 的关键宏，并理解睡眠前后外设时钟处理的职责。"
    ])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FreeRTOS 笔记重点整理")
    set_run_font(fr, size=9, color="777777")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
