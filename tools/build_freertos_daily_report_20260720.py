from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\freertos_study\FreeRTOS今日学习日报-2026-07-20.docx")


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_indent(table, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_indent(table, 120)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    set_font(
        r,
        size={1: 16, 2: 13, 3: 12}.get(level, 11),
        bold=True,
        color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}.get(level, "1F4D78"),
    )
    set_spacing(
        p,
        before={1: 16, 2: 12, 3: 8}.get(level, 6),
        after={1: 8, 2: 6, 3: 4}.get(level, 4),
    )


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    set_spacing(p)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=10)
            set_spacing(p, after=3)
    fix_table_geometry(table, widths)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("FreeRTOS 今日学习日报")
    set_font(tr, size=24, bold=True, color="0B2545")
    set_spacing(title, after=3)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(r"日期：2026年7月20日    代码文件：D:\freertos_study\main.c")
    set_font(sr, size=10, color="555555")
    set_spacing(subtitle, after=16)

    add_heading(doc, "一、今日学习主题", 1)
    add_para(
        doc,
        "今天是在前几天任务创建、挂起恢复、时间管理、信号量和队列集 demo 的基础上继续扩展 FreeRTOS 综合练习。"
        "新增内容重点转向任务通知、软件定时器、Tickless 低功耗和内存管理观察，把笔记中的 API 概念落到可运行、可按键触发、可看日志的代码实验里。",
    )
    add_para(
        doc,
        "代码中新增了第 6 组任务通知日志和第 7 组低功耗日志，并扩展控制台按键：n 用于通知计数，w 用于通知位等待，其他可打印字符用于通知值传递，l 用于切换低功耗教学路径。"
        "这样今天的学习不只是阅读 API，而是通过实际输入、任务阻塞唤醒、日志统计和 hook 函数观察 FreeRTOS 机制如何工作。",
    )

    add_heading(doc, "二、今日新增功能概览", 1)
    add_bullets(
        doc,
        [
            "任务通知：新增三个演示任务，分别覆盖 xTaskNotifyGive + ulTaskNotifyTake 的计数模式、xTaskNotify(..., eSetBits) + xTaskNotifyWait 的事件位模式，以及 eSetValueWithOverwrite 的数值传递模式。",
            "软件定时器：继续使用 heartbeat 周期软件定时器，定时器到期回调只设置 EVENT_HEARTBEAT 事件位，用来体现软件定时器回调运行在 timer service task 中且应保持短小。",
            "低功耗：新增低功耗监视任务、idle hook 计数、tickless 进入前处理 hook 和教学版 suppress hook，可通过 l/L 动态允许或阻止 tickless 路径。",
            "内存管理：在统计任务和启动任务中打印 free heap，并提供 malloc failed hook，体现动态创建任务、队列、信号量、事件组、软件定时器等内核对象时都需要关注 FreeRTOS heap。",
            "控制台实验入口：菜单和按键逻辑把人工输入转化为任务通知和低功耗开关，使今天新增知识点都能通过运行日志直接观察。",
        ],
    )

    add_heading(doc, "三、知识点在代码中的应用", 1)
    add_heading(doc, "1. 任务通知：轻量级的一对一同步与数据通道", 2)
    add_para(
        doc,
        "笔记中提到，每个任务控制块内部都有任务通知值，因此任务通知不需要额外创建队列、信号量或事件组对象，效率高、内存开销小。"
        "今天代码把这个特点拆成三个实验：n/N 按键调用 xTaskNotifyGive()，接收任务用 ulTaskNotifyTake(pdFALSE, portMAX_DELAY) 每次只减 1，效果类似轻量级计数信号量；w/W 按键使用 eSetBits 设置通知位，接收任务用 xTaskNotifyWait() 读取 bit，效果接近单任务事件位；其他可打印按键使用 eSetValueWithOverwrite，把 ASCII 码作为 32 位通知值传给接收任务。",
    )
    add_para(
        doc,
        "通过这三个分支可以看出，任务通知不是单一用法，而是可以按“计数、置位、传值”三种思路组织。"
        "同时也验证了笔记中的限制：任务通知天然面向一个目标任务，不适合广播给多个任务；如果需要多消费者或复杂缓存，仍应选择事件组、队列等对象。",
    )

    add_heading(doc, "2. 软件定时器：定时逻辑不直接放进任务循环", 2)
    add_para(
        doc,
        "软件定时器笔记强调，软件定时器由 FreeRTOS 的 timer service task 管理，启动、停止等操作通过定时器命令队列完成，回调函数不是中断服务函数，但也不应长时间阻塞。"
        "当前代码中的 heartbeat timer 是周期定时器，到期后执行 vHeartbeatTimerCallback()，回调里只做一件事：设置 EVENT_HEARTBEAT 事件位。",
    )
    add_para(
        doc,
        "这种写法把“周期到期”与“业务处理”解耦：定时器只负责产生节拍事件，真正的日志观察或状态处理交给其他任务。"
        "这符合笔记中对软件定时器回调的要求，也能避免把复杂逻辑堆在 timer service task 上影响其他定时器。",
    )

    add_heading(doc, "3. Tickless 低功耗：从 idle 时间观察到睡眠路径控制", 2)
    add_para(
        doc,
        "低功耗笔记中提到，Tickless idle 的核心思想是在系统预计会空闲较长时间时暂停周期 tick，让 MCU 进入睡眠，并在中断唤醒后补偿 tick。"
        "今天代码用 vApplicationIdleHook() 累计空闲 hook 调用次数，用 vLowPowerMonitorTask 周期打印 idle_hook_delta、sleep_prevented_delta、sleep_attempt_delta 和 expected idle ticks，帮助观察系统是否真的有空闲窗口，以及 tickless 路径是否被允许进入。",
    )
    add_para(
        doc,
        "vLowPowerPreSuppressTicksAndSleep() 对应 Tickless 进入前处理：当低功耗 demo 关闭时，把 expected idle time 改为 0，等价于告诉内核本次不要进入低功耗路径。"
        "vLowPowerSuppressTicksAndSleep() 则作为教学版执行 hook，记录预计空闲 tick 和进入次数。Windows 模拟端不会产生真实电流变化，但这段代码把“进入前判断、允许/阻止、进入统计”这条低功耗链路展示出来了。",
    )

    add_heading(doc, "4. 内存管理：从 free heap 到失败 hook", 2)
    add_para(
        doc,
        "内存管理笔记提到，FreeRTOS 动态创建任务、队列、信号量、事件组和软件定时器时，底层会通过 FreeRTOS heap 分配内存；不同 heap_x.c 策略决定是否支持释放、是否合并碎片、是否支持多段内存。"
        "今天代码没有直接写 pvPortMalloc()/vPortFree() 实验，而是把内存观察放进系统运行过程：统计任务周期打印 xPortGetFreeHeapSize()，启动任务在创建全部对象后打印 free_heap，malloc failed hook 则在分配失败时停机等待排查。",
    )
    add_para(
        doc,
        "这使内存管理从“API 名称”变成“创建对象后的资源余量”。随着今天新增三个通知任务和一个低功耗监视任务，free heap 的输出可以帮助判断任务栈和内核对象是否让堆空间明显下降；后续如果加入 pvPortMalloc()/vPortFree() 专项实验，可以继续观察 heap_4 等策略下的释放与碎片合并效果。",
    )

    add_heading(doc, "5. 控制台按键：把知识点变成可重复实验", 2)
    add_para(
        doc,
        "今天的代码把多个知识点都挂到控制台按键上：n 触发通知计数，w 触发通知 bit，普通字符触发通知值传输，l 切换低功耗路径。"
        "这种设计的好处是每个知识点都有明确的输入、阻塞点、唤醒点和日志输出，后续复习时可以按键复现实验，而不是只靠阅读注释理解。",
    )

    add_heading(doc, "四、代码截图行号索引", 1)
    add_para(doc, r"日报正文不直接粘贴代码。需要截图时，可在 D:\freertos_study\main.c 中按下面行号范围截取。")
    add_table(
        doc,
        ["学习内容", "建议截图行号", "截图说明"],
        [
            ["任务通知与低功耗新增资源", "第 51-61 行、第 113-125 行", "展示 LOG_GROUP_6/7、通知 bit、任务通知 demo 句柄和低功耗统计变量。"],
            ["菜单与日志组入口", "第 286-302 行、第 714-719 行", "展示第 6/7 组日志开关，以及运行时查看任务通知和低功耗日志的入口。"],
            ["低功耗开关按键", "第 776-783 行", "展示 l/L 按键如何切换 xLowPowerDemoEnabled，并打印当前 low-power 模式。"],
            ["通知计数发送端", "第 787-796 行", "展示 n/N 按键调用 xTaskNotifyGive()，向 notify_take 任务发送一个计数。"],
            ["通知位发送端", "第 799-810 行", "展示 w/W 按键通过 xTaskNotify(..., eSetBits) 设置 NOTIFY_WAIT_DEMO_BIT。"],
            ["通知值发送端", "第 842-856 行", "展示未占用可打印按键通过 eSetValueWithOverwrite 传递 ASCII 值。"],
            ["软件定时器回调", "第 877-881 行", "展示 vHeartbeatTimerCallback() 运行在 timer service task 中，并只设置 EVENT_HEARTBEAT。"],
            ["任务通知计数接收任务", "第 1062-1095 行", "展示 ulTaskNotifyTake(pdFALSE) 如何逐个消费通知计数。"],
            ["任务通知位等待任务", "第 1109-1142 行", "展示 xTaskNotifyWait() 等待通知 bit 并读取通知值。"],
            ["任务通知值接收任务", "第 1152-1181 行", "展示 xTaskNotifyWait() 读取 ASCII 通知值。"],
            ["低功耗监视任务", "第 1195-1233 行", "展示 idle hook、prevented、attempt、expected idle ticks 的周期统计输出。"],
            ["Tickless 前置与执行 hook", "第 1239-1265 行", "展示进入低功耗前如何阻止 tickless，以及教学版 suppress hook 如何记录预计空闲 tick。"],
            ["新增任务创建", "第 1603-1644 行", "展示三个任务通知 demo 任务和低功耗监视任务的创建。"],
            ["软件定时器创建与启动", "第 1684-1688 行、第 1712 行", "展示 xTimerCreate() 创建周期定时器，并通过 xTimerStart() 启动。"],
            ["内存管理观察与失败 hook", "第 515-518 行、第 1650-1651 行、第 1738-1741 行", "展示 free_heap 输出和 vApplicationMallocFailedHook()。"],
        ],
        [1.55, 1.70, 3.25],
    )

    add_heading(doc, "五、今日收获", 1)
    add_numbered(
        doc,
        [
            "理解了任务通知可以用同一个通知值表达计数、事件位和普通数值三种语义，并能用不同 API 组合实现不同同步模式。",
            "理解了软件定时器回调运行在 timer service task 中，适合做短小触发动作，不适合放长时间阻塞或复杂业务。",
            "通过 idle hook、pre suppress hook 和 suppress hook，建立了 Tickless 低功耗从“预计空闲”到“进入/阻止睡眠路径”的完整观察链路。",
            "把内存管理和对象创建联系起来，知道新增任务、队列、信号量、事件组、软件定时器都会影响 FreeRTOS heap，需要用 free heap 和 malloc failed hook 做运行期观察。",
            "控制台按键让今天的知识点都可以被重复触发，方便对照日志确认任务阻塞、唤醒、通知值变化和低功耗模式切换。",
        ],
    )

    add_heading(doc, "六、问题与后续改进", 1)
    add_bullets(
        doc,
        [
            "当前低功耗 demo 是 Windows 模拟端的教学记录，不会体现真实 MCU 的电流下降；移植到 MCU 时还需要处理外设关闭、SysTick 停止、WFI/WFE、唤醒源和 tick 补偿。",
            "任务通知适合一对一场景，不适合广播给多个任务；如果后续需要多个任务同时感知同一事件，应继续使用事件组或队列等对象。",
            "eSetValueWithOverwrite 会覆盖旧通知值，连续快速按多个普通字符时，接收任务可能只看到最后一个值；如果需要保留每个按键，应该改用队列。",
            "软件定时器回调目前足够短小；如果后续 heartbeat 逻辑变复杂，建议回调只发通知或事件，把复杂处理交给普通任务。",
            "内存管理目前主要是 free heap 观察和 malloc failed hook，后续可以增加 pvPortMalloc()/vPortFree() 小实验，对比 heap_1、heap_4、heap_5 的行为差异。",
            "第 349 行附近仍有采样周期注释与实际延时不一致的问题：代码为 pdMS_TO_TICKS(3000)，注释仍描述为 1000ms/1 秒，建议修正。",
        ],
    )

    add_heading(doc, "七、明日计划", 1)
    add_bullets(
        doc,
        [
            "运行 n、w、普通字符和 l/L 按键实验，记录任务通知三种模式和低功耗开关前后的日志差异。",
            "检查 FreeRTOSConfig.h 中与任务通知、软件定时器、Tickless、malloc failed hook 相关的配置项，和今天代码逐一对应。",
            "补充一个 pvPortMalloc()/vPortFree() 受控实验，观察分配前后 free heap 的变化，并和 heap_4 的碎片合并机制联系起来。",
            "在真实 MCU 或更贴近硬件的平台上继续验证 Tickless idle，观察进入睡眠、唤醒和 tick 补偿的实际效果。",
        ],
    )

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FreeRTOS 今日学习日报 - 2026年7月20日")
    set_font(fr, size=9, color="777777")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
