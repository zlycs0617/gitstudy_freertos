from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\freertos_study\FreeRTOS今日学习日报-2026-07-16.docx")


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
             color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}.get(level, "1F4D78"))
    set_spacing(p, before={1: 16, 2: 12, 3: 8}.get(level, 6),
                after={1: 8, 2: 6, 3: 4}.get(level, 4))


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    set_spacing(p)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=10)
            set_spacing(p, after=3)
    fix_table_geometry(table, widths)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("FreeRTOS 今日学习日报")
    set_font(tr, size=24, bold=True, color="0B2545")
    set_spacing(title, after=3)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("日期：2026年7月16日    代码文件：D:\\freertos_study\\main.c")
    set_font(sr, size=10, color="555555")
    set_spacing(subtitle, after=16)

    add_heading(doc, "一、今日学习主题", 1)
    add_para(doc, "今天是在昨天温湿度监控 demo 的基础上继续扩展 FreeRTOS 练习，重点从“对象间通信”推进到“运行时控制、任务挂起恢复、任务状态查询和系统统计”。新增代码让 demo 不再只是后台自动运行，而是可以通过控制台命令切换日志、手动挂起/恢复任务，并周期性观察任务 API 的真实返回结果。")
    add_para(doc, "这一天的学习更贴近调试阶段：不仅要会创建任务、队列和事件组，还要知道如何观察任务状态、调整任务优先级、查看任务列表、统计运行时间，并通过挂起/恢复验证任务状态转换。")

    add_heading(doc, "二、今日新增功能概览", 1)
    add_bullets(doc, [
        "新增日志分组与控制台命令任务：用 LOG_GROUP_1 到 LOG_GROUP_4 给不同模块日志分类，通过键盘命令实时开关输出。",
        "新增自动与手动两种挂起/恢复演示：vSuspendDemoTask 周期性挂起/恢复 vSuspendWorkerTask，控制台命令也可以手动接管。",
        "新增任务 API 观察模块：vTaskApiWorkerTask 作为被观察对象，vTaskApiInspectorTask 周期性查询状态、修改优先级并打印任务列表和运行时间统计。",
        "新增任务状态名称转换函数，把 eTaskState 枚举转成可读字符串，方便日报和日志中分析 Ready、Blocked、Suspended 等状态。",
        "扩展 Start 任务中的任务创建流程，把新增的控制台任务、挂起演示任务、任务 API 演示任务纳入统一调度。"
    ])

    add_heading(doc, "三、知识点在代码中的应用", 1)
    add_heading(doc, "1. 日志分组：从单纯打印到可控观察", 2)
    add_para(doc, "昨天的日志输出已经用互斥量保护，解决了多任务同时打印时互相穿插的问题。今天进一步增加日志分组掩码，把统计/参数、传感器/控制、监督、挂起 worker 四类输出分别映射到不同 bit。这样做把“事件位/位掩码”的思想迁移到了调试工具本身：通过一个整数的不同 bit 控制不同模块是否输出。")
    add_para(doc, "控制台命令任务使用 Windows 桌面环境下的 _kbhit/_getch 轮询键盘输入，再调用切换函数改变日志掩码。虽然这不是 FreeRTOS 通用外设中断写法，但它适合当前桌面 port 学习：可以在系统运行中实时观察某一组任务，而不必修改代码重新编译。")

    add_heading(doc, "2. 挂起与恢复：验证任务状态转换", 2)
    add_para(doc, "今天新增的 vSuspendWorkerTask 是一个持续打印运行计数和 tick 的工作任务，vSuspendDemoTask 则周期性调用 vTaskSuspend 和 vTaskResume。运行时可以直接看到 [4] 组日志停止和恢复，这把笔记中“挂起态不参与调度，恢复后重新进入就绪态”的概念变成了可观察现象。")
    add_para(doc, "控制台命令还提供手动挂起/恢复，并用 xManualSuspendWorker 防止自动恢复覆盖手动状态。这个细节体现了状态管理意识：当系统同时存在自动控制和人工控制时，需要有一个标志位表达当前控制权，避免两个控制源互相打架。")

    add_heading(doc, "3. 任务查询 API：从任务会跑到任务可诊断", 2)
    add_para(doc, "今天新增的 task api worker/inspector 模块集中练习了任务相关 API。worker 自己调用 uxTaskPriorityGet(NULL) 打印当前优先级；inspector 通过 xTaskGetHandle 按任务名找到 worker，再通过 uxTaskPriorityGet、vTaskPrioritySet、uxTaskGetNumberOfTasks、eTaskGetState、uxTaskGetStackHighWaterMark 等 API 观察和改变任务。")
    add_para(doc, "这部分对应笔记中“任务相关的其他 API 函数”章节。学习重点不再只是知道 API 名称，而是理解它们各自解决的问题：句柄获取、优先级调整、任务数量统计、任务状态判断、栈水位评估，都是调试 FreeRTOS 系统时会用到的工具。")

    add_heading(doc, "4. 任务快照与运行时间统计", 2)
    add_para(doc, "vTaskGetInfo 用于获取单个任务的详细快照，uxTaskGetSystemState 用于一次性获取系统任务数组，vTaskList 和 vTaskGetRunTimeStats 则生成调试用文本表格。今天把这些接口集中放在 inspector 任务里周期性输出，说明已经开始从“局部任务逻辑”转向“全系统视角”。")
    add_para(doc, "vTaskList 和 vTaskGetRunTimeStats 对 FreeRTOSConfig.h 有配置依赖，也需要比较大的输出缓冲区。代码中用静态缓冲区保存表格输出，避免把大数组放到任务栈上，这和栈水位观察结合起来，是一个很好的工程习惯。")

    add_heading(doc, "5. 优先级动态调整与调度观察", 2)
    add_para(doc, "inspector 周期性把 worker 的优先级在 tskIDLE_PRIORITY + 1 和 tskIDLE_PRIORITY + 3 之间切换。这样日志中可以直接看到优先级改变，并进一步观察任务状态、执行频率和调度表现。这个练习把抢占式调度的知识点推进到了运行时实验：优先级不是固定概念，而是可以通过 API 动态改变的调度参数。")

    add_heading(doc, "6. 任务创建规模扩大后的系统组织", 2)
    add_para(doc, "Start 任务现在除了创建昨天已有的 sensor、control、supervisor、stats、param 等任务，还创建了 suspend worker、suspend demo、console command、task api worker、task api inspector。任务数增加后，统一在 Start 任务临界段内创建，并用 configASSERT 检查创建结果，仍然保持了清晰的初始化路径。")

    add_heading(doc, "四、代码截图行号索引", 1)
    add_para(doc, "日报正文不直接粘贴代码。需要截图时，可在 D:\\freertos_study\\main.c 中按下面行号范围截取。")
    add_table(doc, ["新增功能", "建议截图行号", "截图说明"], [
        ["控制台输入支持", "第 3 行", "引入 conio.h，用于 Windows 桌面 demo 的 _kbhit/_getch 实时命令输入。"],
        ["日志分组位定义", "第 45-49 行", "用 LOG_GROUP_1 到 LOG_GROUP_4 和 LOG_GROUP_ALL 定义不同日志组。"],
        ["新增任务句柄和控制标志", "第 51-52 行、第 90-95 行", "保存 task api 演示任务、挂起演示任务、控制台命令任务句柄，以及日志掩码/手动挂起标志。"],
        ["日志过滤与互斥打印", "第 113-163 行", "在 prvPrintLine 中按日志组过滤，再用 xConsoleMutex 保护输出；另有块输出和日志菜单函数。"],
        ["自动挂起/恢复演示", "第 491-538 行", "vSuspendWorkerTask 持续运行，vSuspendDemoTask 自动调用 vTaskSuspend/vTaskResume。"],
        ["控制台命令任务", "第 543-613 行", "用键盘命令切换日志组，手动挂起/恢复 worker，并打印帮助菜单。"],
        ["任务状态名称转换", "第 642-669 行", "把 eTaskState 枚举转换为 Running、Ready、Blocked、Suspended 等可读文本。"],
        ["任务 API worker", "第 672-703 行", "worker 打印自身优先级，并主动延时让 inspector 可以观察 Blocked 状态。"],
        ["任务 API inspector", "第 717-839 行", "集中调用任务查询、优先级修改、系统快照、任务列表和运行时间统计 API。"],
        ["新增任务创建", "第 900-939 行", "Start 任务中创建挂起演示、控制台命令和 task api 演示相关任务。"]
    ], [1.65, 1.55, 3.30])

    add_heading(doc, "五、今日收获", 1)
    add_numbered(doc, [
        "通过挂起/恢复演示，把“挂起态不参与调度、恢复后重新进入就绪态”的笔记内容变成了运行现象。",
        "通过日志分组和控制台命令，理解了调试输出也需要设计，否则任务数量增多后日志会淹没关键信息。",
        "通过任务 API inspector，系统性练习了句柄获取、优先级查询/修改、任务数量、任务状态、栈水位、任务快照和运行时间统计。",
        "开始认识到 FreeRTOS 学习不只是会用通信对象，还要会观察系统、定位任务状态和评估资源占用。",
        "通过静态缓冲区保存任务列表和运行时间统计输出，进一步理解了任务栈空间管理的重要性。"
    ])

    add_heading(doc, "六、问题与后续改进", 1)
    add_bullets(doc, [
        "第 263 行附近仍有注释与实际延时不一致的问题：代码是 pdMS_TO_TICKS(3000)，注释仍描述为 1000ms/1 秒，后续应统一。",
        "当前控制台命令使用轮询方式，每 100ms 检查一次键盘输入；后续可以思考如果是真实外设输入，应该如何改成中断或队列事件驱动。",
        "vTaskList 和 vTaskGetRunTimeStats 适合学习和调试，但正式项目中要注意开销和配置依赖，后续可对比 uxTaskGetSystemState 的结构化方式。",
        "任务优先级动态调整已经能观察效果，下一步可以记录调整前后 worker 日志频率和状态变化，形成更明确的实验结论。"
    ])

    add_heading(doc, "七、明日计划", 1)
    add_bullets(doc, [
        "整理 FreeRTOSConfig.h 中与任务查询、运行时间统计、任务列表、栈溢出检测相关的宏配置。",
        "补一组实验日志：手动挂起 worker、自动恢复被手动接管、再手动恢复，验证状态机是否符合预期。",
        "尝试把控制台命令产生的动作封装成队列消息，由专门控制任务消费，进一步练习队列在控制面中的使用。",
        "修正采样周期注释，并检查新增代码中的注释是否都和实际参数一致。"
    ])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FreeRTOS 今日学习日报 - 2026年7月16日")
    set_font(fr, size=9, color="777777")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
