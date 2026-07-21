from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\freertos_study\FreeRTOS今日学习日报-2026-07-17.docx")


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
             color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}.get(level, "1F4D78"))
    set_spacing(p, before={1: 16, 2: 12, 3: 8}.get(level, 6),
                after={1: 8, 2: 6, 3: 4}.get(level, 4))


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    set_spacing(p)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)
        set_spacing(p, after=5, line=1.167)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=10)
            set_spacing(p, after=3)
    fix_table_geometry(table, widths)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("FreeRTOS 今日学习日报")
    set_font(tr, size=24, bold=True, color="0B2545")
    set_spacing(title, after=3)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("日期：2026年7月17日    代码文件：D:\\freertos_study\\main.c")
    set_font(sr, size=10, color="555555")
    set_spacing(subtitle, after=16)

    add_heading(doc, "一、今日学习主题", 1)
    add_para(doc, "今天是在前两天综合 demo 的基础上继续扩展 FreeRTOS 练习，重点从任务观察和挂起恢复推进到时间管理、信号量同步和队列集。新增代码把第 5 组日志作为“时间/信号量/队列集实验区”，通过控制台按键触发二值信号量、计数型信号量和队列消息，再由队列集监视任务统一等待和处理。")
    add_para(doc, "这一天的学习更强调“同一任务如何等待多种事件源”。相对延时和绝对延时用于对比任务周期控制；二值信号量和计数信号量用于对比同步事件是否可累计；队列集则把队列和信号量组合成一个统一等待入口。")

    add_heading(doc, "二、今日新增功能概览", 1)
    add_bullets(doc, [
        "新增 LOG_GROUP_5，把延时、信号量和队列集相关日志单独分组，便于运行时开关观察。",
        "新增相对延时演示任务 vRelativeDelayDemoTask，用 vTaskDelay() 展示“从调用时刻开始延时”的特点。",
        "新增绝对延时演示任务 vAbsoluteDelayDemoTask，用 vTaskDelayUntil() 展示“按固定周期唤醒”的特点。",
        "新增二值信号量和计数型信号量，分别通过控制台按键 b/c 释放，观察二值事件和可累计事件的差异。",
        "新增队列集 xDemoQueueSet，把二值信号量、计数信号量和消息队列加入同一个等待集合，由 vQueueSetMonitorTask 统一阻塞等待。",
        "扩展控制台命令任务，新增 b、c、q 三类输入，把人工按键转化为 FreeRTOS 同步对象和队列消息。"
    ])

    add_heading(doc, "三、知识点在代码中的应用", 1)
    add_heading(doc, "1. 时间管理：相对延时与绝对延时对照", 2)
    add_para(doc, "笔记中提到，vTaskDelay() 是相对延时，每次都从调用函数的时刻开始计算；vTaskDelayUntil() 是绝对延时，更适合周期任务。今天代码用两个独立任务做对照：相对延时任务在延时前后打印 tick，绝对延时任务使用 xLastWakeTime 作为周期基准，让任务按固定节拍唤醒。")
    add_para(doc, "这个练习把“延时函数”从概念变成了可观察实验。相对延时适合简单让出 CPU；绝对延时更适合采样、控制、心跳这类需要稳定周期的任务。")

    add_heading(doc, "2. 二值信号量：一次事件只保留一个 token", 2)
    add_para(doc, "今天新增的 xBinaryDemoSemaphore 用于演示二值信号量。控制台按 b/B 时调用 xSemaphoreGive，如果信号量已经处于可用状态，再次释放不会继续累计，日志会提示 binary semaphore already full。")
    add_para(doc, "这对应笔记里“二值信号量只有空和满两种情况”的描述。它适合表达一次同步事件是否发生，但不适合统计连续发生了多少次事件。")

    add_heading(doc, "3. 计数型信号量：事件可以累计", 2)
    add_para(doc, "计数信号量 xCountingDemoSemaphore 初始计数为 0，最大计数为 SEM_DEMO_COUNTING_MAX。控制台按 c/C 时释放一次信号量，计数值增加；队列集监视任务处理时再 xSemaphoreTake，把计数值减一。")
    add_para(doc, "这对应笔记里“事件计数”的使用场景：事件发生时 give，处理任务 take。和二值信号量相比，计数型信号量能保留多次事件，因此适合按键次数、资源数量、待处理事件数量这类场景。")

    add_heading(doc, "4. 队列集：一个任务等待多类对象", 2)
    add_para(doc, "今天最核心的新功能是队列集。xDemoQueueSet 同时加入了二值信号量、计数信号量和一个 uint32_t 消息队列。vQueueSetMonitorTask 不再分别阻塞等待三个对象，而是只阻塞在 xQueueSelectFromSet() 上，哪个成员先就绪就返回哪个成员。")
    add_para(doc, "队列集只告诉任务“哪个成员可读/可取”，真正消费还要对具体成员调用 xSemaphoreTake() 或 xQueueReceive()。这个细节很重要：选择成员和消费成员是两步，不能只调用 xQueueSelectFromSet() 就认为事件已经处理完成。")

    add_heading(doc, "5. 控制台输入变成事件源", 2)
    add_para(doc, "控制台命令任务今天不只是切换日志组和手动挂起恢复，还新增了 b、c、q 三类输入。b 释放二值信号量，c 释放计数信号量，q 发送队列消息。这样人工按键就变成了不同类型的事件源，统一交给队列集监视任务处理。")
    add_para(doc, "虽然当前使用 _kbhit/_getch 是 Windows 桌面 demo 的便利方式，但设计思想可以迁移到嵌入式场景：外设输入或中断事件最终也可以转化为信号量、队列消息或事件位，再由任务集中处理。")

    add_heading(doc, "6. 初始化顺序与队列集容量", 2)
    add_para(doc, "main() 中先创建二值信号量、计数信号量、队列和队列集，再把这些成员加入队列集，最后启动任务和调度器。队列集容量设置为 1 + SEM_DEMO_COUNTING_MAX + QUEUE_SET_DEMO_QUEUE_LENGTH，体现了容量要能覆盖成员可能产生的事件数量。")
    add_para(doc, "这和队列集笔记中的注意点一致：队列或信号量应先创建并保持空状态，再加入队列集；队列集长度要根据成员容量综合估算。")

    add_heading(doc, "四、代码截图行号索引", 1)
    add_para(doc, "日报正文不直接粘贴代码。需要截图时，可在 D:\\freertos_study\\main.c 中按下面行号范围截取。")
    add_table(doc, ["新增功能", "建议截图行号", "截图说明"], [
        ["第 5 组日志与容量宏", "第 49-55 行", "展示 LOG_GROUP_5、计数信号量最大计数、队列集队列长度等新增宏。"],
        ["新增任务句柄与同步对象", "第 100-106 行", "展示相对/绝对延时任务、队列集监视任务、二值/计数信号量、队列和队列集句柄。"],
        ["菜单新增按键说明", "第 252-254 行", "展示 b、c、q 等按键如何对应信号量和队列集 demo。"],
        ["控制台触发同步对象", "第 629-741 行", "展示 b 释放二值信号量、c 释放计数信号量、q 发送队列消息。"],
        ["相对延时任务", "第 776-805 行", "展示 vTaskDelay() 前后打印 tick，观察相对延时从调用点开始计算。"],
        ["绝对延时任务", "第 813-850 行", "展示 xLastWakeTime 和 vTaskDelayUntil()，观察固定周期唤醒。"],
        ["队列集监视任务", "第 864-929 行", "展示 xQueueSelectFromSet() 等待多个成员，并分别 take/receive 消费。"],
        ["新增任务创建", "第 1237-1267 行", "展示 Start 任务中创建相对延时、绝对延时和队列集监视任务。"],
        ["信号量和队列集创建", "第 1293-1321 行", "展示创建二值信号量、计数信号量、队列集队列、队列集，并把成员加入队列集。"]
    ], [1.65, 1.55, 3.30])

    add_heading(doc, "五、今日收获", 1)
    add_numbered(doc, [
        "通过两个延时任务，对 vTaskDelay() 和 vTaskDelayUntil() 的差异有了运行层面的认识。",
        "通过按键 b/c 的对照，理解了二值信号量不累计、计数型信号量可累计的核心区别。",
        "通过队列集监视任务，理解了一个任务可以用 xQueueSelectFromSet() 等待多个队列/信号量成员。",
        "理解了队列集返回的是“已就绪成员”，真正消费还要再调用 xSemaphoreTake() 或 xQueueReceive()。",
        "进一步体会到调试 demo 的可操作性很重要：控制台按键让同步对象的状态变化更容易观察。"
    ])

    add_heading(doc, "六、问题与后续改进", 1)
    add_bullets(doc, [
        "第 349 行附近仍有采样周期注释与实际延时不一致的问题：代码是 pdMS_TO_TICKS(3000)，注释仍描述为 1000ms/1 秒，建议尽快修正。",
        "当前控制台按键直接 give/send，同步对象变化比较直观；后续可以把按键输入先封装成控制消息，再由一个控制任务统一处理。",
        "队列集当前包含二值信号量、计数信号量和普通队列，下一步可以尝试移除成员或动态调整成员，验证队列集成员必须为空才能移除的规则。",
        "绝对延时任务可以加入一段模拟工作耗时，对比任务执行时间接近周期时，vTaskDelayUntil() 的表现。"
    ])

    add_heading(doc, "七、明日计划", 1)
    add_bullets(doc, [
        "结合日志实际运行一次 b/c/q 输入实验，记录队列集返回成员和消费顺序。",
        "补充一组 vTaskDelay 与 vTaskDelayUntil 的 tick 输出对比，验证周期稳定性。",
        "继续阅读队列集和信号量笔记，把 API 参数、返回值和适用场景补成个人速查表。",
        "检查 FreeRTOSConfig.h 中与队列集、计数信号量、软件定时器和任务统计相关的配置项。"
    ])

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FreeRTOS 今日学习日报 - 2026年7月17日")
    set_font(fr, size=9, color="777777")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
